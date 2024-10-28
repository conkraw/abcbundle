import streamlit as st
from docx import Document
from io import BytesIO
from datetime import datetime
import firebase_admin
from firebase_admin import credentials, firestore
import os
import json
from datetime import datetime
import io
import pytz 

# Define mappings for ETT size, Blade type, and Apneic Oxygenation based on patient age
age_to_ett_mapping = {'': '', 
                      '0 months': '3.5 mm',
 '1 months': '4.0 mm',
 '2 months': '4.0 mm',
 '3 months': '4.0 mm',
 '4 months': '4.0 mm',
 '5 months': '4.0 mm',
 '6 months': '4.0 mm',
 '7 months': '4.0 mm',
 '8 months': '4.0 mm',
 '9 months': '4.0 mm',
 '10 months': '4.0 mm',
 '11 months': '4.0 mm',
 '12 months': '4.0 mm',
 '1 year': '3.5 mm',
 '2 years': '4.0 mm',
 '3 years': '4.0 mm',
 '4 years': '4.5 mm',
 '5 years': '4.5 mm',
 '6 years': '5.0 mm',
 '7 years': '5.0 mm',
 '8 years': '5.5 mm',
 '9 years': '5.5 mm',
 '10 years': '6.0 mm',
 '11 years': '6.0 mm',
 '12 years': '6.5 mm',
 '13 years': '6.5 mm',
 '14 years': '7.0 mm',
 '15 years': '7.0 mm',
 '16 years': '7.5 mm',
 '17 years': '7.5 mm',
 '18 years': '7.5 mm',
 '19 years': '7.5 mm',
 '20 years': '7.5 mm',
 '21 years': '7.5 mm',
 '22 years': '7.5 mm',
 '23 years': '7.5 mm',
 '24 years': '7.5 mm',
 '25 years': '7.5 mm'}

lma_size_mapping = {
    '2.0 kg': 'LMA Size 1',
    '2.5 kg': 'LMA Size 1',
    '3.0 kg': 'LMA Size 1',
    '3.5 kg': 'LMA Size 1',
    '4.0 kg': 'LMA Size 1',
    '4.5 kg': 'LMA Size 1',
    '5.0 kg': 'LMA Size 1.5',
    '5.5 kg': 'LMA Size 1.5',
    '6.0 kg': 'LMA Size 1.5',
    '6.5 kg': 'LMA Size 1.5',
    '7.0 kg': 'LMA Size 1.5',
    '7.5 kg': 'LMA Size 1.5',
    '8.0 kg': 'LMA Size 1.5',
    '8.5 kg': 'LMA Size 1.5',
    '9.0 kg': 'LMA Size 1.5',
    '9.5 kg': 'LMA Size 1.5',
    '10.0 kg': 'LMA Size 2',
    '11.0 kg': 'LMA Size 2',
    '12.0 kg': 'LMA Size 2',
    '13.0 kg': 'LMA Size 2',
    '14.0 kg': 'LMA Size 2',
    '15.0 kg': 'LMA Size 2',
    '16.0 kg': 'LMA Size 2',
    '17.0 kg': 'LMA Size 2',
    '18.0 kg': 'LMA Size 2',
    '19.0 kg': 'LMA Size 2',
    '20.0 kg': 'LMA Size 2.5',
    '21.0 kg': 'LMA Size 2.5',
    '22.0 kg': 'LMA Size 2.5',
    '23.0 kg': 'LMA Size 2.5',
    '24.0 kg': 'LMA Size 2.5',
    '25.0 kg': 'LMA Size 2.5',
    '26.0 kg': 'LMA Size 3',
    '27.0 kg': 'LMA Size 3',
    '28.0 kg': 'LMA Size 3',
    '29.0 kg': 'LMA Size 3',
    '30.0 kg': 'LMA Size 3',
    '31.0 kg': 'LMA Size 3',
    '32.0 kg': 'LMA Size 3',
    '33.0 kg': 'LMA Size 3',
    '34.0 kg': 'LMA Size 3',
    '35.0 kg': 'LMA Size 3',
    '36.0 kg': 'LMA Size 3',
    '37.0 kg': 'LMA Size 3',
    '38.0 kg': 'LMA Size 3',
    '39.0 kg': 'LMA Size 3',
    '40.0 kg': 'LMA Size 3',
    '41.0 kg': 'LMA Size 4',
    '42.0 kg': 'LMA Size 4',
    '43.0 kg': 'LMA Size 4',
    '44.0 kg': 'LMA Size 4',
    '45.0 kg': 'LMA Size 4',
    '46.0 kg': 'LMA Size 4',
    '47.0 kg': 'LMA Size 4',
    '48.0 kg': 'LMA Size 4',
    '49.0 kg': 'LMA Size 4',
    '50.0 kg': 'LMA Size 4',
    '51.0 kg': 'LMA Size 5',
    '52.0 kg': 'LMA Size 5',
    '53.0 kg': 'LMA Size 5',
    '54.0 kg': 'LMA Size 5',
    '55.0 kg': 'LMA Size 5',
    '56.0 kg': 'LMA Size 5',
    '57.0 kg': 'LMA Size 5',
    '58.0 kg': 'LMA Size 5',
    '59.0 kg': 'LMA Size 5',
    '60.0 kg': 'LMA Size 5',
    '61.0 kg': 'LMA Size 5',
    '62.0 kg': 'LMA Size 5',
    '63.0 kg': 'LMA Size 5',
    '64.0 kg': 'LMA Size 5',
    '65.0 kg': 'LMA Size 5',
    '66.0 kg': 'LMA Size 5',
    '67.0 kg': 'LMA Size 5',
    '68.0 kg': 'LMA Size 5',
    '69.0 kg': 'LMA Size 5',
    '70.0 kg': 'LMA Size 5',
    '71.0 kg': 'LMA Size 5',
    '72.0 kg': 'LMA Size 5',
    '73.0 kg': 'LMA Size 5',
    '74.0 kg': 'LMA Size 5',
    '75.0 kg': 'LMA Size 5',
    '76.0 kg': 'LMA Size 5',
    '77.0 kg': 'LMA Size 5',
    '78.0 kg': 'LMA Size 5',
    '79.0 kg': 'LMA Size 5',
    '80.0 kg': 'LMA Size 5',
    '81.0 kg': 'LMA Size 5',
    '82.0 kg': 'LMA Size 5',
    '83.0 kg': 'LMA Size 5',
    '84.0 kg': 'LMA Size 5',
    '85.0 kg': 'LMA Size 5',
    '86.0 kg': 'LMA Size 5',
    '87.0 kg': 'LMA Size 5',
    '88.0 kg': 'LMA Size 5',
    '89.0 kg': 'LMA Size 5',
    '90.0 kg': 'LMA Size 5',
    '91.0 kg': 'LMA Size 5',
    '92.0 kg': 'LMA Size 5',
    '93.0 kg': 'LMA Size 5',
    '94.0 kg': 'LMA Size 5',
    '95.0 kg': 'LMA Size 5',
    '96.0 kg': 'LMA Size 5',
    '97.0 kg': 'LMA Size 5',
    '98.0 kg': 'LMA Size 5',
    '99.0 kg': 'LMA Size 5',
    '100.0 kg': 'LMA Size 5'
}


age_to_glide_mapping = {
    '': '', 
    '0 months': 'Glidescope Size 1 or 2',
    '1 months': 'Glidescope Size 1 or 2',
    '2 months': 'Glidescope Size 1 or 2',
    '3 months': 'Glidescope Size 1 or 2',
    '4 months': 'Glidescope Size 1 or 2',
    '5 months': 'Glidescope Size 1 or 2',
    '6 months': 'Glidescope Size 2',
    '7 months': 'Glidescope Size 2',
    '8 months': 'Glidescope Size 2',
    '9 months': 'Glidescope Size 2',
    '10 months': 'Glidescope Size 2',
    '11 months': 'Glidescope Size 2',
    '12 months': 'Glidescope Size 2',
    '1 year': 'Glidescope Size 2',
    '2 years': 'Glidescope Size 3',
    '3 years': 'Glidescope Size 3',
    '4 years': 'Glidescope Size 3',
    '5 years': 'Glidescope Size 3',
    '6 years': 'Glidescope Size 3',
    '7 years': 'Glidescope Size 3',
    '8 years': 'Glidescope Size 3',
    '9 years': 'Glidescope Size 3',
    '10 years': 'Glidescope Size 3',
    '11 years': 'Glidescope Size 3 or 4',
    '12 years': 'Glidescope Size 3 or 4',
    '13 years': 'Glidescope Size 3 or 4',
    '14 years': 'Glidescope Size 3 or 4',
    '15 years': 'Glidescope Size 3 or 4',
    '16 years': 'Glidescope Size 3 or 4',
    '17 years': 'Glidescope Size 3 or 4',
    '18 years': 'Glidescope Size 3 or 4',
    '19 years': 'Glidescope Size 3 or 4',
    '20 years': 'Glidescope Size 3 or 4',
    '21 years': 'Glidescope Size 3 or 4',
    '22 years': 'Glidescope Size 3 or 4',
    '23 years': 'Glidescope Size 3 or 4',
    '24 years': 'Glidescope Size 3 or 4',
    '25 years': 'Glidescope Size 3 or 4'
}


# Define other mappings based on age
age_to_mac_mapping = {
    '': '', 
    '0 months': '',
    '1 months': '',
    '2 months': '',
    '3 months': '',
    '4 months': '',
    '5 months': '',
    '6 months': '',
    '7 months': '',
    '8 months': '',
    '9 months': '',
    '10 months': '',
    '11 months': '',
    '12 months': 'Mac Blade Size 2',
    '1 year': 'Mac Blade Size 2',
    '2 years': 'Mac Blade Size 2',
    '3 years': 'Mac Blade Size 2',
    '4 years': 'Mac Blade Size 2',
    '5 years': 'Mac Blade Size 2',
    '6 years': 'Mac Blade Size 2',
    '7 years': 'Mac Blade Size 2',
    '8 years': 'Mac Blade Size 2',
    '9 years': 'Mac Blade Size 2',
    '10 years': 'Mac Blade Size 2',
    '11 years': 'Mac Blade Size 3',
    '12 years': 'Mac Blade Size 3',
    '13 years': 'Mac Blade Size 3',
    '14 years': 'Mac Blade Size 3',
    '15 years': 'Mac Blade Size 3',
    '16 years': 'Mac Blade Size 3',
    '17 years': 'Mac Blade Size 3',
    '18 years': 'Mac Blade Size 3',
    '19 years': 'Mac Blade Size 3',
    '20 years': 'Mac Blade Size 3',
    '21 years': 'Mac Blade Size 3',
    '22 years': 'Mac Blade Size 3',
    '23 years': 'Mac Blade Size 3',
    '24 years': 'Mac Blade Size 3',
    '25 years': 'Mac Blade Size 3'
}

age_to_miller_mapping = {
    '': '', 
    '0 months': 'Miller Blade Size 0',
    '1 months': 'Miller Blade Size 0',
    '2 months': 'Miller Blade Size 0',
    '3 months': 'Miller Blade Size 0',
    '4 months': 'Miller Blade Size 0',
    '5 months': 'Miller Blade Size 0',
    '6 months': 'Miller Blade Size 1',
    '7 months': 'Miller Blade Size 1',
    '8 months': 'Miller Blade Size 1',
    '9 months': 'Miller Blade Size 1',
    '10 months': 'Miller Blade Size 1',
    '11 months': 'Miller Blade Size 1',
    '12 months': 'Miller Blade Size 1',
    '1 year': 'Miller Blade Size 1',
    '2 years': 'Miller Blade Size 1',
    '3 years': 'Miller Blade Size 1',
    '4 years': 'Miller Blade Size 2',
    '5 years': 'Miller Blade Size 2',
    '6 years': 'Miller Blade Size 2',
    '7 years': 'Miller Blade Size 2',
    '8 years': 'Miller Blade Size 2',
    '9 years': 'Miller Blade Size 2',
    '10 years': 'Miller Blade Size 2',
    '11 years': 'Miller Blade Size 2',
    '12 years': 'Miller Blade Size 2',
    '13 years': 'Miller Blade Size 2',
    '14 years': 'Miller Blade Size 2',
    '15 years': 'Miller Blade Size 2',
    '16 years': 'Miller Blade Size 2',
    '17 years': 'Miller Blade Size 2',
    '18 years': 'Miller Blade Size 2',
    '19 years': 'Miller Blade Size 2',
    '20 years': 'Miller Blade Size 2',
    '21 years': 'Miller Blade Size 2',
    '22 years': 'Miller Blade Size 2',
    '23 years': 'Miller Blade Size 2',
    '24 years': 'Miller Blade Size 2',
    '25 years': 'Miller Blade Size 2'
}


weight_to_atropine_mapping = {'2.0 kg': '0.1 mg', '2.5 kg': '0.1 mg', '3.0 kg': '0.1 mg', '3.5 kg': '0.1 mg', '4.0 kg': '0.1 mg', '4.5 kg': '0.1 mg', '5.0 kg': '0.1 mg', '5.5 kg': '0.11 mg', '6.0 kg': '0.12 mg', '6.5 kg': '0.13 mg', '7.0 kg': '0.14 mg', '7.5 kg': '0.15 mg', '8.0 kg': '0.16 mg', '8.5 kg': '0.17 mg', '9.0 kg': '0.18 mg', '9.5 kg': '0.19 mg', '10.0 kg': '0.2 mg', '11.0 kg': '0.22 mg', '12.0 kg': '0.24 mg', '13.0 kg': '0.26 mg', '14.0 kg': '0.28 mg', '15.0 kg': '0.3 mg', '16.0 kg': '0.32 mg', '17.0 kg': '0.34 mg', '18.0 kg': '0.36 mg', '19.0 kg': '0.38 mg', '20.0 kg': '0.4 mg', '21.0 kg': '0.42 mg', '22.0 kg': '0.44 mg', '23.0 kg': '0.46 mg', '24.0 kg': '0.48 mg', '25.0 kg': '0.5 mg', '26.0 kg': '0.52 mg', '27.0 kg': '0.54 mg', '28.0 kg': '0.56 mg', '29.0 kg': '0.58 mg', '30.0 kg': '0.6 mg', '31.0 kg': '0.62 mg', '32.0 kg': '0.64 mg', '33.0 kg': '0.66 mg', '34.0 kg': '0.68 mg', '35.0 kg': '0.7 mg', '36.0 kg': '0.72 mg', '37.0 kg': '0.74 mg', '38.0 kg': '0.76 mg', '39.0 kg': '0.78 mg', '40.0 kg': '0.8 mg', '41.0 kg': '0.82 mg', '42.0 kg': '0.84 mg', '43.0 kg': '0.86 mg', '44.0 kg': '0.88 mg', '45.0 kg': '0.9 mg', '46.0 kg': '0.92 mg', '47.0 kg': '0.94 mg', '48.0 kg': '0.96 mg', '49.0 kg': '0.98 mg', '50.0 kg': '1.0 mg', '51.0 kg': '1.0 mg', '52.0 kg': '1.0 mg', '53.0 kg': '1.0 mg', '54.0 kg': '1.0 mg', '55.0 kg': '1.0 mg', '56.0 kg': '1.0 mg', '57.0 kg': '1.0 mg', '58.0 kg': '1.0 mg', '59.0 kg': '1.0 mg', '60.0 kg': '1.0 mg', '61.0 kg': '1.0 mg', '62.0 kg': '1.0 mg', '63.0 kg': '1.0 mg', '64.0 kg': '1.0 mg', '65.0 kg': '1.0 mg', '66.0 kg': '1.0 mg', '67.0 kg': '1.0 mg', '68.0 kg': '1.0 mg', '69.0 kg': '1.0 mg', '70.0 kg': '1.0 mg', '71.0 kg': '1.0 mg', '72.0 kg': '1.0 mg', '73.0 kg': '1.0 mg', '74.0 kg': '1.0 mg', '75.0 kg': '1.0 mg', '76.0 kg': '1.0 mg', '77.0 kg': '1.0 mg', '78.0 kg': '1.0 mg', '79.0 kg': '1.0 mg', '80.0 kg': '1.0 mg', '81.0 kg': '1.0 mg', '82.0 kg': '1.0 mg', '83.0 kg': '1.0 mg', '84.0 kg': '1.0 mg', '85.0 kg': '1.0 mg', '86.0 kg': '1.0 mg', '87.0 kg': '1.0 mg', '88.0 kg': '1.0 mg', '89.0 kg': '1.0 mg', '90.0 kg': '1.0 mg', '91.0 kg': '1.0 mg', '92.0 kg': '1.0 mg', '93.0 kg': '1.0 mg', '94.0 kg': '1.0 mg', '95.0 kg': '1.0 mg', '96.0 kg': '1.0 mg', '97.0 kg': '1.0 mg', '98.0 kg': '1.0 mg', '99.0 kg': '1.0 mg', '100.0 kg': '1.0 mg'}

weight_to_glycopyrrolate_mapping = {'2.0 kg': '0.008 mg', '2.5 kg': '0.012 mg', '3.0 kg': '0.012 mg', '3.5 kg': '0.016 mg', '4.0 kg': '0.016 mg', '4.5 kg': '0.02 mg', '5.0 kg': '0.02 mg', '5.5 kg': '0.024 mg', '6.0 kg': '0.024 mg', '6.5 kg': '0.028 mg', '7.0 kg': '0.028 mg', '7.5 kg': '0.032 mg', '8.0 kg': '0.032 mg', '8.5 kg': '0.036 mg', '9.0 kg': '0.036 mg', '9.5 kg': '0.04 mg', '10.0 kg': '0.04 mg', '11.0 kg': '0.048 mg', '12.0 kg': '0.048 mg', '13.0 kg': '0.056 mg', '14.0 kg': '0.056 mg', '15.0 kg': '0.064 mg', '16.0 kg': '0.064 mg', '17.0 kg': '0.072 mg', '18.0 kg': '0.072 mg', '19.0 kg': '0.08 mg', '20.0 kg': '0.08 mg', '21.0 kg': '0.08 mg', '22.0 kg': '0.08 mg', '23.0 kg': '0.1 mg', '24.0 kg': '0.1 mg', '25.0 kg': '0.1 mg', '26.0 kg': '0.1 mg', '27.0 kg': '0.1 mg', '28.0 kg': '0.1 mg', '29.0 kg': '0.1 mg', '30.0 kg': '0.1 mg', '31.0 kg': '0.1 mg', '32.0 kg': '0.1 mg', '33.0 kg': '0.1 mg', '34.0 kg': '0.1 mg', '35.0 kg': '0.1 mg', '36.0 kg': '0.1 mg', '37.0 kg': '0.1 mg', '38.0 kg': '0.1 mg', '39.0 kg': '0.1 mg', '40.0 kg': '0.1 mg', '41.0 kg': '0.1 mg', '42.0 kg': '0.1 mg', '43.0 kg': '0.1 mg', '44.0 kg': '0.1 mg', '45.0 kg': '0.1 mg', '46.0 kg': '0.1 mg', '47.0 kg': '0.1 mg', '48.0 kg': '0.1 mg', '49.0 kg': '0.1 mg', '50.0 kg': '0.1 mg', '51.0 kg': '0.1 mg', '52.0 kg': '0.1 mg', '53.0 kg': '0.1 mg', '54.0 kg': '0.1 mg', '55.0 kg': '0.1 mg', '56.0 kg': '0.1 mg', '57.0 kg': '0.1 mg', '58.0 kg': '0.1 mg', '59.0 kg': '0.1 mg', '60.0 kg': '0.1 mg', '61.0 kg': '0.1 mg', '62.0 kg': '0.1 mg', '63.0 kg': '0.1 mg', '64.0 kg': '0.1 mg', '65.0 kg': '0.1 mg', '66.0 kg': '0.1 mg', '67.0 kg': '0.1 mg', '68.0 kg': '0.1 mg', '69.0 kg': '0.1 mg', '70.0 kg': '0.1 mg', '71.0 kg': '0.1 mg', '72.0 kg': '0.1 mg', '73.0 kg': '0.1 mg', '74.0 kg': '0.1 mg', '75.0 kg': '0.1 mg', '76.0 kg': '0.1 mg', '77.0 kg': '0.1 mg', '78.0 kg': '0.1 mg', '79.0 kg': '0.1 mg', '80.0 kg': '0.1 mg', '81.0 kg': '0.1 mg', '82.0 kg': '0.1 mg', '83.0 kg': '0.1 mg', '84.0 kg': '0.1 mg', '85.0 kg': '0.1 mg', '86.0 kg': '0.1 mg', '87.0 kg': '0.1 mg', '88.0 kg': '0.1 mg', '89.0 kg': '0.1 mg', '90.0 kg': '0.1 mg', '91.0 kg': '0.1 mg', '92.0 kg': '0.1 mg', '93.0 kg': '0.1 mg', '94.0 kg': '0.1 mg', '95.0 kg': '0.1 mg', '96.0 kg': '0.1 mg', '97.0 kg': '0.1 mg', '98.0 kg': '0.1 mg', '99.0 kg': '0.1 mg', '100.0 kg': '0.1 mg'}

weight_to_fentanyl_mapping = {'2.0 kg': '2.0 mcg', '2.5 kg': '3.0 mcg', '3.0 kg': '3.0 mcg', '3.5 kg': '4.0 mcg', '4.0 kg': '4.0 mcg', '4.5 kg': '5.0 mcg', '5.0 kg': '5.0 mcg', '5.5 kg': '6.0 mcg', '6.0 kg': '6.0 mcg', '6.5 kg': '7.0 mcg', '7.0 kg': '7.0 mcg', '7.5 kg': '8.0 mcg', '8.0 kg': '8.0 mcg', '8.5 kg': '9.0 mcg', '9.0 kg': '9.0 mcg', '9.5 kg': '10.0 mcg', '10.0 kg': '10.0 mcg', '11.0 kg': '12.0 mcg', '12.0 kg': '12.0 mcg', '13.0 kg': '14.0 mcg', '14.0 kg': '14.0 mcg', '15.0 kg': '16.0 mcg', '16.0 kg': '16.0 mcg', '17.0 kg': '18.0 mcg', '18.0 kg': '18.0 mcg', '19.0 kg': '20.0 mcg', '20.0 kg': '20.0 mcg', '21.0 kg': '20.0 mcg', '22.0 kg': '20.0 mcg', '23.0 kg': '25.0 mcg', '24.0 kg': '25.0 mcg', '25.0 kg': '25.0 mcg', '26.0 kg': '25.0 mcg', '27.0 kg': '25.0 mcg', '28.0 kg': '30.0 mcg', '29.0 kg': '30.0 mcg', '30.0 kg': '30.0 mcg', '31.0 kg': '30.0 mcg', '32.0 kg': '30.0 mcg', '33.0 kg': '35.0 mcg', '34.0 kg': '35.0 mcg', '35.0 kg': '35.0 mcg', '36.0 kg': '35.0 mcg', '37.0 kg': '35.0 mcg', '38.0 kg': '40.0 mcg', '39.0 kg': '40.0 mcg', '40.0 kg': '40.0 mcg', '41.0 kg': '40.0 mcg', '42.0 kg': '40.0 mcg', '43.0 kg': '45.0 mcg', '44.0 kg': '45.0 mcg', '45.0 kg': '45.0 mcg', '46.0 kg': '45.0 mcg', '47.0 kg': '45.0 mcg', '48.0 kg': '45.0 mcg', '49.0 kg': '45.0 mcg', '50.0 kg': '50.0 mcg', '51.0 kg': '50.0 mcg', '52.0 kg': '50.0 mcg', '53.0 kg': '50.0 mcg', '54.0 kg': '50.0 mcg', '55.0 kg': '50.0 mcg', '56.0 kg': '50.0 mcg', '57.0 kg': '50.0 mcg', '58.0 kg': '50.0 mcg', '59.0 kg': '50.0 mcg', '60.0 kg': '50.0 mcg', '61.0 kg': '50.0 mcg', '62.0 kg': '50.0 mcg', '63.0 kg': '50.0 mcg', '64.0 kg': '50.0 mcg', '65.0 kg': '50.0 mcg', '66.0 kg': '50.0 mcg', '67.0 kg': '50.0 mcg', '68.0 kg': '50.0 mcg', '69.0 kg': '50.0 mcg', '70.0 kg': '50.0 mcg', '71.0 kg': '50.0 mcg', '72.0 kg': '50.0 mcg', '73.0 kg': '50.0 mcg', '74.0 kg': '50.0 mcg', '75.0 kg': '50.0 mcg', '76.0 kg': '50.0 mcg', '77.0 kg': '50.0 mcg', '78.0 kg': '50.0 mcg', '79.0 kg': '50.0 mcg', '80.0 kg': '50.0 mcg', '81.0 kg': '50.0 mcg', '82.0 kg': '50.0 mcg', '83.0 kg': '50.0 mcg', '84.0 kg': '50.0 mcg', '85.0 kg': '50.0 mcg', '86.0 kg': '50.0 mcg', '87.0 kg': '50.0 mcg', '88.0 kg': '50.0 mcg', '89.0 kg': '50.0 mcg', '90.0 kg': '50.0 mcg', '91.0 kg': '50.0 mcg', '92.0 kg': '50.0 mcg', '93.0 kg': '50.0 mcg', '94.0 kg': '50.0 mcg', '95.0 kg': '50.0 mcg', '96.0 kg': '50.0 mcg', '97.0 kg': '50.0 mcg', '98.0 kg': '50.0 mcg', '99.0 kg': '50.0 mcg', '100.0 kg': '50.0 mcg'}

weight_to_midaz_mapping = {'2.0 kg': '0.1 mg', '2.5 kg': '0.15 mg', '3.0 kg': '0.15 mg', '3.5 kg': '0.2 mg', '4.0 kg': '0.2 mg', '4.5 kg': '0.25 mg', '5.0 kg': '0.25 mg', '5.5 kg': '0.3 mg', '6.0 kg': '0.3 mg', '6.5 kg': '0.35 mg', '7.0 kg': '0.35 mg', '7.5 kg': '0.4 mg', '8.0 kg': '0.4 mg', '8.5 kg': '0.45 mg', '9.0 kg': '0.45 mg', '9.5 kg': '0.5 mg', '10.0 kg': '0.5 mg', '11.0 kg': '0.6 mg', '12.0 kg': '0.6 mg', '13.0 kg': '0.7 mg', '14.0 kg': '0.7 mg', '15.0 kg': '0.8 mg', '16.0 kg': '0.8 mg', '17.0 kg': '0.9 mg', '18.0 kg': '0.9 mg', '19.0 kg': '1.0 mg', '20.0 kg': '1.0 mg', '21.0 kg': '1.0 mg', '22.0 kg': '1.0 mg', '23.0 kg': '1.25 mg', '24.0 kg': '1.25 mg', '25.0 kg': '1.25 mg', '26.0 kg': '1.25 mg', '27.0 kg': '1.25 mg', '28.0 kg': '1.5 mg', '29.0 kg': '1.5 mg', '30.0 kg': '1.5 mg', '31.0 kg': '1.5 mg', '32.0 kg': '1.5 mg', '33.0 kg': '1.5 mg', '34.0 kg': '1.75 mg', '35.0 kg': '1.75 mg', '36.0 kg': '1.75 mg', '37.0 kg': '1.75 mg', '38.0 kg': '2.0 mg', '39.0 kg': '2.0 mg', '40.0 kg': '2.0 mg', '41.0 kg': '2.0 mg', '42.0 kg': '2.0 mg', '43.0 kg': '2.0 mg', '44.0 kg': '2.0 mg', '45.0 kg': '2.0 mg', '46.0 kg': '2.0 mg', '47.0 kg': '2.0 mg', '48.0 kg': '2.0 mg', '49.0 kg': '2.0 mg', '50.0 kg': '2.0 mg', '51.0 kg': '2.0 mg', '52.0 kg': '2.0 mg', '53.0 kg': '2.0 mg', '54.0 kg': '2.0 mg', '55.0 kg': '2.0 mg', '56.0 kg': '2.0 mg', '57.0 kg': '2.0 mg', '58.0 kg': '2.0 mg', '59.0 kg': '2.0 mg', '60.0 kg': '2.0 mg', '61.0 kg': '2.0 mg', '62.0 kg': '2.0 mg', '63.0 kg': '2.0 mg', '64.0 kg': '2.0 mg', '65.0 kg': '2.0 mg', '66.0 kg': '2.0 mg', '67.0 kg': '2.0 mg', '68.0 kg': '2.0 mg', '69.0 kg': '2.0 mg', '70.0 kg': '2.0 mg', '71.0 kg': '2.0 mg', '72.0 kg': '2.0 mg', '73.0 kg': '2.0 mg', '74.0 kg': '2.0 mg', '75.0 kg': '2.0 mg', '76.0 kg': '2.0 mg', '77.0 kg': '2.0 mg', '78.0 kg': '2.0 mg', '79.0 kg': '2.0 mg', '80.0 kg': '2.0 mg', '81.0 kg': '2.0 mg', '82.0 kg': '2.0 mg', '83.0 kg': '2.0 mg', '84.0 kg': '2.0 mg', '85.0 kg': '2.0 mg', '86.0 kg': '2.0 mg', '87.0 kg': '2.0 mg', '88.0 kg': '2.0 mg', '89.0 kg': '2.0 mg', '90.0 kg': '2.0 mg', '91.0 kg': '2.0 mg', '92.0 kg': '2.0 mg', '93.0 kg': '2.0 mg', '94.0 kg': '2.0 mg', '95.0 kg': '2.0 mg', '96.0 kg': '2.0 mg', '97.0 kg': '2.0 mg', '98.0 kg': '2.0 mg', '99.0 kg': '2.0 mg', '100.0 kg': '2.0 mg'}

weight_to_ketamine_mapping = {'2.0 kg': '0.1 mg', '2.5 kg': '0.15 mg', '3.0 kg': '0.15 mg', '3.5 kg': '0.2 mg', '4.0 kg': '0.2 mg', '4.5 kg': '0.25 mg', '5.0 kg': '0.25 mg', '5.5 kg': '0.3 mg', '6.0 kg': '0.3 mg', '6.5 kg': '0.35 mg', '7.0 kg': '0.35 mg', '7.5 kg': '0.4 mg', '8.0 kg': '0.4 mg', '8.5 kg': '0.45 mg', '9.0 kg': '0.45 mg', '9.5 kg': '0.5 mg', '10.0 kg': '0.5 mg', '11.0 kg': '0.6 mg', '12.0 kg': '0.6 mg', '13.0 kg': '0.7 mg', '14.0 kg': '0.7 mg', '15.0 kg': '0.8 mg', '16.0 kg': '0.8 mg', '17.0 kg': '0.9 mg', '18.0 kg': '0.9 mg', '19.0 kg': '1.0 mg', '20.0 kg': '1.0 mg', '21.0 kg': '1.0 mg', '22.0 kg': '1.0 mg', '23.0 kg': '1.25 mg', '24.0 kg': '1.25 mg', '25.0 kg': '1.25 mg', '26.0 kg': '1.25 mg', '27.0 kg': '1.25 mg', '28.0 kg': '1.5 mg', '29.0 kg': '1.5 mg', '30.0 kg': '1.5 mg', '31.0 kg': '1.5 mg', '32.0 kg': '1.5 mg', '33.0 kg': '1.5 mg', '34.0 kg': '1.75 mg', '35.0 kg': '1.75 mg', '36.0 kg': '1.75 mg', '37.0 kg': '1.75 mg', '38.0 kg': '2.0 mg', '39.0 kg': '2.0 mg', '40.0 kg': '2.0 mg', '41.0 kg': '2.0 mg', '42.0 kg': '2.0 mg', '43.0 kg': '2.0 mg', '44.0 kg': '2.0 mg', '45.0 kg': '2.0 mg', '46.0 kg': '2.0 mg', '47.0 kg': '2.0 mg', '48.0 kg': '2.0 mg', '49.0 kg': '2.0 mg', '50.0 kg': '2.0 mg', '51.0 kg': '2.0 mg', '52.0 kg': '2.0 mg', '53.0 kg': '2.0 mg', '54.0 kg': '2.0 mg', '55.0 kg': '2.0 mg', '56.0 kg': '2.0 mg', '57.0 kg': '2.0 mg', '58.0 kg': '2.0 mg', '59.0 kg': '2.0 mg', '60.0 kg': '2.0 mg', '61.0 kg': '2.0 mg', '62.0 kg': '2.0 mg', '63.0 kg': '2.0 mg', '64.0 kg': '2.0 mg', '65.0 kg': '2.0 mg', '66.0 kg': '2.0 mg', '67.0 kg': '2.0 mg', '68.0 kg': '2.0 mg', '69.0 kg': '2.0 mg', '70.0 kg': '2.0 mg', '71.0 kg': '2.0 mg', '72.0 kg': '2.0 mg', '73.0 kg': '2.0 mg', '74.0 kg': '2.0 mg', '75.0 kg': '2.0 mg', '76.0 kg': '2.0 mg', '77.0 kg': '2.0 mg', '78.0 kg': '2.0 mg', '79.0 kg': '2.0 mg', '80.0 kg': '2.0 mg', '81.0 kg': '2.0 mg', '82.0 kg': '2.0 mg', '83.0 kg': '2.0 mg', '84.0 kg': '2.0 mg', '85.0 kg': '2.0 mg', '86.0 kg': '2.0 mg', '87.0 kg': '2.0 mg', '88.0 kg': '2.0 mg', '89.0 kg': '2.0 mg', '90.0 kg': '2.0 mg', '91.0 kg': '2.0 mg', '92.0 kg': '2.0 mg', '93.0 kg': '2.0 mg', '94.0 kg': '2.0 mg', '95.0 kg': '2.0 mg', '96.0 kg': '2.0 mg', '97.0 kg': '2.0 mg', '98.0 kg': '2.0 mg', '99.0 kg': '2.0 mg', '100.0 kg': '2.0 mg'}

weight_to_propo_mapping = {'2.0 kg': '0.1 mg', '2.5 kg': '0.15 mg', '3.0 kg': '0.15 mg', '3.5 kg': '0.2 mg', '4.0 kg': '0.2 mg', '4.5 kg': '0.25 mg', '5.0 kg': '0.25 mg', '5.5 kg': '0.3 mg', '6.0 kg': '0.3 mg', '6.5 kg': '0.35 mg', '7.0 kg': '0.35 mg', '7.5 kg': '0.4 mg', '8.0 kg': '0.4 mg', '8.5 kg': '0.45 mg', '9.0 kg': '0.45 mg', '9.5 kg': '0.5 mg', '10.0 kg': '0.5 mg', '11.0 kg': '0.6 mg', '12.0 kg': '0.6 mg', '13.0 kg': '0.7 mg', '14.0 kg': '0.7 mg', '15.0 kg': '0.8 mg', '16.0 kg': '0.8 mg', '17.0 kg': '0.9 mg', '18.0 kg': '0.9 mg', '19.0 kg': '1.0 mg', '20.0 kg': '1.0 mg', '21.0 kg': '1.0 mg', '22.0 kg': '1.0 mg', '23.0 kg': '1.25 mg', '24.0 kg': '1.25 mg', '25.0 kg': '1.25 mg', '26.0 kg': '1.25 mg', '27.0 kg': '1.25 mg', '28.0 kg': '1.5 mg', '29.0 kg': '1.5 mg', '30.0 kg': '1.5 mg', '31.0 kg': '1.5 mg', '32.0 kg': '1.5 mg', '33.0 kg': '1.5 mg', '34.0 kg': '1.75 mg', '35.0 kg': '1.75 mg', '36.0 kg': '1.75 mg', '37.0 kg': '1.75 mg', '38.0 kg': '2.0 mg', '39.0 kg': '2.0 mg', '40.0 kg': '2.0 mg', '41.0 kg': '2.0 mg', '42.0 kg': '2.0 mg', '43.0 kg': '2.0 mg', '44.0 kg': '2.0 mg', '45.0 kg': '2.0 mg', '46.0 kg': '2.0 mg', '47.0 kg': '2.0 mg', '48.0 kg': '2.0 mg', '49.0 kg': '2.0 mg', '50.0 kg': '2.0 mg', '51.0 kg': '2.0 mg', '52.0 kg': '2.0 mg', '53.0 kg': '2.0 mg', '54.0 kg': '2.0 mg', '55.0 kg': '2.0 mg', '56.0 kg': '2.0 mg', '57.0 kg': '2.0 mg', '58.0 kg': '2.0 mg', '59.0 kg': '2.0 mg', '60.0 kg': '2.0 mg', '61.0 kg': '2.0 mg', '62.0 kg': '2.0 mg', '63.0 kg': '2.0 mg', '64.0 kg': '2.0 mg', '65.0 kg': '2.0 mg', '66.0 kg': '2.0 mg', '67.0 kg': '2.0 mg', '68.0 kg': '2.0 mg', '69.0 kg': '2.0 mg', '70.0 kg': '2.0 mg', '71.0 kg': '2.0 mg', '72.0 kg': '2.0 mg', '73.0 kg': '2.0 mg', '74.0 kg': '2.0 mg', '75.0 kg': '2.0 mg', '76.0 kg': '2.0 mg', '77.0 kg': '2.0 mg', '78.0 kg': '2.0 mg', '79.0 kg': '2.0 mg', '80.0 kg': '2.0 mg', '81.0 kg': '2.0 mg', '82.0 kg': '2.0 mg', '83.0 kg': '2.0 mg', '84.0 kg': '2.0 mg', '85.0 kg': '2.0 mg', '86.0 kg': '2.0 mg', '87.0 kg': '2.0 mg', '88.0 kg': '2.0 mg', '89.0 kg': '2.0 mg', '90.0 kg': '2.0 mg', '91.0 kg': '2.0 mg', '92.0 kg': '2.0 mg', '93.0 kg': '2.0 mg', '94.0 kg': '2.0 mg', '95.0 kg': '2.0 mg', '96.0 kg': '2.0 mg', '97.0 kg': '2.0 mg', '98.0 kg': '2.0 mg', '99.0 kg': '2.0 mg', '100.0 kg': '2.0 mg'}

weight_to_roc_mapping = {'2.0 kg': '2.0 mg', '2.5 kg': '2.5 mg', '3.0 kg': '3.0 mg', '3.5 kg': '3.5 mg', '4.0 kg': '4.0 mg', '4.5 kg': '4.5 mg', '5.0 kg': '5.0 mg', '5.5 kg': '5.5 mg', '6.0 kg': '6.0 mg', '6.5 kg': '6.5 mg', '7.0 kg': '7.0 mg', '7.5 kg': '7.5 mg', '8.0 kg': '8.0 mg', '8.5 kg': '8.5 mg', '9.0 kg': '9.0 mg', '9.5 kg': '9.5 mg', '10.0 kg': '10.0 mg', '11.0 kg': '11.0 mg', '12.0 kg': '12.0 mg', '13.0 kg': '13.0 mg', '14.0 kg': '14.0 mg', '15.0 kg': '15.0 mg', '16.0 kg': '16.0 mg', '17.0 kg': '17.0 mg', '18.0 kg': '18.0 mg', '19.0 kg': '19.0 mg', '20.0 kg': '20.0 mg', '21.0 kg': '21.0 mg', '22.0 kg': '22.0 mg', '23.0 kg': '23.0 mg', '24.0 kg': '24.0 mg', '25.0 kg': '25.0 mg', '26.0 kg': '26.0 mg', '27.0 kg': '27.0 mg', '28.0 kg': '28.0 mg', '29.0 kg': '29.0 mg', '30.0 kg': '30.0 mg', '31.0 kg': '31.0 mg', '32.0 kg': '32.0 mg', '33.0 kg': '33.0 mg', '34.0 kg': '34.0 mg', '35.0 kg': '35.0 mg', '36.0 kg': '36.0 mg', '37.0 kg': '37.0 mg', '38.0 kg': '38.0 mg', '39.0 kg': '39.0 mg', '40.0 kg': '40.0 mg', '41.0 kg': '41.0 mg', '42.0 kg': '42.0 mg', '43.0 kg': '43.0 mg', '44.0 kg': '44.0 mg', '45.0 kg': '45.0 mg', '46.0 kg': '46.0 mg', '47.0 kg': '47.0 mg', '48.0 kg': '48.0 mg', '49.0 kg': '49.0 mg', '50.0 kg': '50.0 mg', '51.0 kg': '50.0 mg', '52.0 kg': '50.0 mg', '53.0 kg': '50.0 mg', '54.0 kg': '50.0 mg', '55.0 kg': '50.0 mg', '56.0 kg': '50.0 mg', '57.0 kg': '50.0 mg', '58.0 kg': '50.0 mg', '59.0 kg': '50.0 mg', '60.0 kg': '50.0 mg', '61.0 kg': '50.0 mg', '62.0 kg': '50.0 mg', '63.0 kg': '50.0 mg', '64.0 kg': '50.0 mg', '65.0 kg': '50.0 mg', '66.0 kg': '50.0 mg', '67.0 kg': '50.0 mg', '68.0 kg': '50.0 mg', '69.0 kg': '50.0 mg', '70.0 kg': '50.0 mg', '71.0 kg': '50.0 mg', '72.0 kg': '50.0 mg', '73.0 kg': '50.0 mg', '74.0 kg': '50.0 mg', '75.0 kg': '50.0 mg', '76.0 kg': '50.0 mg', '77.0 kg': '50.0 mg', '78.0 kg': '50.0 mg', '79.0 kg': '50.0 mg', '80.0 kg': '50.0 mg', '81.0 kg': '50.0 mg', '82.0 kg': '50.0 mg', '83.0 kg': '50.0 mg', '84.0 kg': '50.0 mg', '85.0 kg': '50.0 mg', '86.0 kg': '50.0 mg', '87.0 kg': '50.0 mg', '88.0 kg': '50.0 mg', '89.0 kg': '50.0 mg', '90.0 kg': '50.0 mg', '91.0 kg': '50.0 mg', '92.0 kg': '50.0 mg', '93.0 kg': '50.0 mg', '94.0 kg': '50.0 mg', '95.0 kg': '50.0 mg', '96.0 kg': '50.0 mg', '97.0 kg': '50.0 mg', '98.0 kg': '50.0 mg', '99.0 kg': '50.0 mg', '100.0 kg': '50.0 mg'}

weight_to_vec_mapping = {'2.0 kg': '0.2 mg', '2.5 kg': '0.3 mg', '3.0 kg': '0.3 mg', '3.5 kg': '0.4 mg', '4.0 kg': '0.4 mg', '4.5 kg': '0.5 mg', '5.0 kg': '0.5 mg', '5.5 kg': '0.6 mg', '6.0 kg': '0.6  mg', '6.5 kg': '0.7 mg', '7.0 kg': '0.7 mg', '7.5 kg': '0.8 mg', '8.0 kg': '0.8 mg', '8.5 kg': '0.9 mg', '9.0 kg': '0.9 mg', '9.5 kg': '1.0 mg', '10.0 kg': '1.0 mg', '11.0 kg': '1.2 mg', '12.0 kg': '1.2 mg', '13.0 kg': '1.4 mg', '14.0 kg': '1.4 mg', '15.0 kg': '1.6 mg', '16.0 kg': '1.6 mg', '17.0 kg': '1.8 mg', '18.0 kg': '1.8 mg', '19.0 kg': '2.0 mg', '20.0 kg': '2.0 mg', '21.0 kg': '2.0 mg', '22.0 kg': '2.0 mg', '23.0 kg': '2.5 mg', '24.0 kg': '2.5 mg', '25.0 kg': '2.5 mg', '26.0 kg': '2.5 mg', '27.0 kg': '2.5 mg', '28.0 kg': '3.0 mg', '29.0 kg': '3.0 mg', '30.0 kg': '3.0 mg', '31.0 kg': '3.0 mg', '32.0 kg': '3.0 mg', '33.0 kg': '3.5 mg', '34.0 kg': '3.5 mg', '35.0 kg': '3.5 mg', '36.0 kg': '3.5 mg', '37.0 kg': '3.5 mg', '38.0 kg': '4.0 mg', '39.0 kg': '4.0 mg', '40.0 kg': '4.0 mg', '41.0 kg': '4.0 mg', '42.0 kg': '4.0 mg', '43.0 kg': '4.5 mg', '44.0 kg': '4.5 mg', '45.0 kg': '4.5 mg', '46.0 kg': '4.5 mg', '47.0 kg': '4.5 mg', '48.0 kg': '4.5 mg', '49.0 kg': '4.5 mg', '50.0 kg': '10.0 mg', '51.0 kg': '10.0 mg', '52.0 kg': '10.0 mg', '53.0 kg': '10.0 mg', '54.0 kg': '10.0 mg', '55.0 kg': '10.0 mg', '56.0 kg': '10.0 mg', '57.0 kg': '10.0 mg', '58.0 kg': '10.0 mg', '59.0 kg': '10.0 mg', '60.0 kg': '10.0 mg', '61.0 kg': '10.0 mg', '62.0 kg': '10.0 mg', '63.0 kg': '10.0 mg', '64.0 kg': '10.0 mg', '65.0 kg': '10.0 mg', '66.0 kg': '10.0 mg', '67.0 kg': '10.0 mg', '68.0 kg': '10.0 mg', '69.0 kg': '10.0 mg', '70.0 kg': '10.0 mg', '71.0 kg': '10.0 mg', '72.0 kg': '10.0 mg', '73.0 kg': '10.0 mg', '74.0 kg': '10.0 mg', '75.0 kg': '10.0 mg', '76.0 kg': '10.0 mg', '77.0 kg': '10.0 mg', '78.0 kg': '10.0 mg', '79.0 kg': '10.0 mg', '80.0 kg': '10.0 mg', '81.0 kg': '10.0 mg', '82.0 kg': '10.0 mg', '83.0 kg': '10.0 mg', '84.0 kg': '10.0 mg', '85.0 kg': '10.0 mg', '86.0 kg': '10.0 mg', '87.0 kg': '10.0 mg', '88.0 kg': '10.0 mg', '89.0 kg': '10.0 mg', '90.0 kg': '10.0 mg', '91.0 kg': '10.0 mg', '92.0 kg': '10.0 mg', '93.0 kg': '10.0 mg', '94.0 kg': '10.0 mg', '95.0 kg': '10.0 mg', '96.0 kg': '10.0 mg', '97.0 kg': '10.0 mg', '98.0 kg': '10.0 mg', '99.0 kg': '10.0 mg', '100.0 kg': '10.0 mg'}

age_to_oxygenation_mapping = {'0 months': '5 Liters Per Minute', '1 months': '5 Liters Per Minute', '2 months': '5 Liters Per Minute', '3 months': '5 Liters Per Minute', '4 months': '5 Liters Per Minute', '5 months': '5 Liters Per Minute', '6 months': '5 Liters Per Minute', '7 months': '5 Liters Per Minute', '8 months': '5 Liters Per Minute', '9 months': '5 Liters Per Minute', '10 months': '5 Liters Per Minute', '11 months': '5 Liters Per Minute', '12 months': '10 Liters Per Minute', '1 year': '10 Liters Per Minute', '2 years': '10 Liters Per Minute', '3 years': '10 Liters Per Minute', '4 years': '10 Liters Per Minute', '5 years': '10 Liters Per Minute', '6 years': '10 Liters Per Minute', '7 years': '10 Liters Per Minute', '8 years': '15 Liters Per Minute', '9 years': '15 Liters Per Minute', '10 years': '15 Liters Per Minute', '11 years': '15 Liters Per Minute', '12 years': '15 Liters Per Minute', '13 years': '15 Liters Per Minute', '14 years': '15 Liters Per Minute', '15 years': '15 Liters Per Minute', '16 years': '15 Liters Per Minute', '17 years': '15 Liters Per Minute', '18 years': '15 Liters Per Minute', '19 years': '15 Liters Per Minute', '20 years': '15 Liters Per Minute', '21 years': '15 Liters Per Minute', '22 years': '15 Liters Per Minute', '23 years': '15 Liters Per Minute', '24 years': '15 Liters Per Minute', '25 years': '15 Liters Per Minute'}

# Define a function to automatically update the other settings when the age is selected
def update_automatic_selections():
    # Check if age is selected (you can keep this or modify it as needed)
    if "age_select" in st.session_state and st.session_state.age_select:
        selected_age = st.session_state.age_select
        st.session_state.ett_size = age_to_ett_mapping[selected_age]
        st.session_state.glide_details = age_to_glide_mapping[selected_age]
        st.session_state.mac_details = age_to_mac_mapping[selected_age]
        st.session_state.miller_details = age_to_miller_mapping[selected_age]
        st.session_state.oxygenation = age_to_oxygenation_mapping[selected_age]

    # Check if weight is selected
    if "weight_select" in st.session_state and st.session_state.weight_select:
        selected_weight = st.session_state.weight_select
        # Update drug dosages based on the selected weight
        st.session_state.lma_details = lma_size_mapping[selected_weight]
        st.session_state.atr_dose = weight_to_atropine_mapping[selected_weight]
        st.session_state.gly_dose = weight_to_glycopyrrolate_mapping[selected_weight]
        st.session_state.fen_dose = weight_to_fentanyl_mapping[selected_weight]
        st.session_state.mid_dose = weight_to_midaz_mapping[selected_weight]
        st.session_state.ket_dose = weight_to_ketamine_mapping[selected_weight]
        st.session_state.pro_dose = weight_to_propo_mapping[selected_weight]
        st.session_state.roc_dose = weight_to_roc_mapping[selected_weight]
        st.session_state.vec_dose = weight_to_vec_mapping[selected_weight]

def create_word_doc(template_path, data):
    # Load the Word document template
    doc = Document(template_path)
    
    # Access parameters
    date = data.get('date')
    time = data.get('time')
    option = data.get('option')
    completed_by = data.get('completed_by')
    room_number = data.get('room_number')
    difficult_airway_history = data.get('difficult_airway_history')
    physical_risk = data.get('physical_risk')
    high_risk_desaturation = data.get('high_risk_desaturation')
    high_risk_ICP = data.get('high_risk_ICP')
    unstable_hemodynamics = data.get('unstable_hemodynamics')
    other_risk_yes_no = data.get('other_risk_yes_no')
    other_risk_text_input = data.get('other_risk_text_input')
    who_will_intubate = data.get('who_will_intubate')
    who_will_bvm = data.get('who_will_bvm')
    other_intubate = data.get('other_intubate')
    other_bvm = data.get('other_bvm')
    intubation_method = data.get('intubation_method')
    ett_size = data.get('ett_size')
    ett_type = data.get('ett_type')
    lma_details = data.get('lma_details')
    glide_details = data.get('glide_details')
    other_device_details = data.get('other_device_details')
    mac_details = data.get('mac_details')
    miller_details = data.get('miller_details')
    wis_hipple_details = data.get('wis_hipple_details')
    atropine_dose = data.get('atropine_dose')
    glycopyrrolate_dose = data.get('glycopyrrolate_dose')
    fentanyl_dose = data.get('fentanyl_dose')
    midazolam_dose = data.get('midazolam_dose')
    ketamine_dose = data.get('ketamine_dose')
    propofol_dose = data.get('propofol_dose')
    roc_dose = data.get('roc_dose')
    vec_dose = data.get('vec_dose')
    ao_details = data.get('ao_details')
    other_planning = data.get('other_planning')
    when_intubate = data.get('when_intubate')
    advance_airway_provider = data.get('advance_airway_provider')
    advance_airway_procedure = data.get('advance_airway_procedure')




    # Check and replace text in paragraphs
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            # Replace Date and Time Placeholders
            if 'DatePlaceholder' in run.text:
                run.text = run.text.replace('DatePlaceholder', date)
            if 'TimePlaceholder' in run.text:
                run.text = run.text.replace('TimePlaceholder', time)
            if 'FrontPagePlaceholder' in run.text:
                run.text = run.text.replace('FrontPagePlaceholder', option)
            if 'DocumenterPlaceholder' in run.text:
                run.text = run.text.replace('DocumenterPlaceholder', completed_by)
            if 'room_number' in run.text:
                run.text = run.text.replace('room_number', room_number)
            if 'D1' in run.text:
                run.text = run.text.replace('D1', difficult_airway_history)
            if 'D2' in run.text:
                run.text = run.text.replace('D2', physical_risk)
            if 'R1' in run.text:
                run.text = run.text.replace('R1', high_risk_desaturation)
            if 'R2' in run.text:
                run.text = run.text.replace('R2', high_risk_ICP)
            if 'R3' in run.text:
                run.text = run.text.replace('R3', unstable_hemodynamics)
            if 'R4' in run.text:
                run.text = run.text.replace('R4', other_risk_yes_no)
            if 'risk_factors' in run.text:
                run.text = run.text.replace('risk_factors', other_risk_text_input)
            if 'who_will_intubate' in run.text:
                if who_will_intubate:
                    run.text = run.text.replace('who_will_intubate', ', '.join(who_will_intubate).rstrip(', '))  # Join with comma and space, then strip
            if 'who_will_bvm' in run.text:
                if who_will_bvm:
                    run.text = run.text.replace('who_will_bvm', ', '.join(who_will_bvm).rstrip(', '))  # Join with comma and space, then strip
            if 'intubation_method' in run.text:
                run.text = run.text.replace('intubation_method', intubation_method)
            if 'ett_type' in run.text:
                run.text = run.text.replace('ett_type', ett_type)
            if 'ett_size' in run.text:
                run.text = run.text.replace('ett_size', ett_size)
            if 'lma_details' in run.text:
                run.text = run.text.replace('lma_details', lma_details)
            if 'glide_details' in run.text:
                run.text = run.text.replace('glide_details', glide_details)
            if 'other_device_details' in run.text:
                run.text = run.text.replace('other_device_details', other_device_details)
            if 'mac_details' in run.text:
                run.text = run.text.replace('mac_details', mac_details)
            if 'miller_details' in run.text:
                run.text = run.text.replace('miller_details', miller_details)
            if 'wis_hipple_details' in run.text:
                run.text = run.text.replace('wis_hipple_details', wis_hipple_details)
            if 'atropine_dose' in run.text:
                run.text = run.text.replace('atropine_dose', atropine_dose)
            if 'glycopyrrolate_dose' in run.text:
                run.text = run.text.replace('glycopyrrolate_dose', glycopyrrolate_dose)
            if 'fentanyl_dose' in run.text:
                run.text = run.text.replace('fentanyl_dose', fentanyl_dose)
            if 'midazolam_dose' in run.text:
                run.text = run.text.replace('midazolam_dose', midazolam_dose)
            if 'ketamine_dose' in run.text:
                run.text = run.text.replace('ketamine_dose', ketamine_dose)
            if 'propofol_dose' in run.text:
                run.text = run.text.replace('propofol_dose', propofol_dose)
            if 'roc_dose' in run.text:
                run.text = run.text.replace('roc_dose', roc_dose)
            if 'vec_dose' in run.text:
                run.text = run.text.replace('vec_dose', vec_dose)
            if 'ao_details' in run.text:
                run.text = run.text.replace('ao_details', ao_details)
            if 'additional_notes' in run.text:
                run.text = run.text.replace('additional_notes', other_planning)
            if 'when_intubate' in run.text:
                if when_intubate:
                    run.text = run.text.replace('when_intubate', ', '.join(when_intubate).rstrip(', '))  # Join with comma and space, then strip
            if 'advance_airway_provider' in run.text:
                if advance_airway_provider:
                    run.text = run.text.replace('advance_airway_provider', ', '.join(advance_airway_provider).rstrip(', '))  # Join with comma and space, then strip
            if 'advance_airway_procedure' in run.text:
                if advance_airway_procedure:
                    run.text = run.text.replace('advance_airway_procedure', ', '.join(advance_airway_procedure).rstrip(', '))  # Join with comma and space, then strip

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        # Replace Date and Time Placeholders
                        if 'DatePlaceholder' in run.text:
                            run.text = run.text.replace('DatePlaceholder', date)
                        if 'TimePlaceholder' in run.text:
                            run.text = run.text.replace('TimePlaceholder', time)
                        if 'FrontPagePlaceholder' in run.text:
                            run.text = run.text.replace('FrontPagePlaceholder', option)
                        if 'DocumenterPlaceholder' in run.text:
                            run.text = run.text.replace('DocumenterPlaceholder', completed_by)
                        if 'room_number' in run.text:
                            run.text = run.text.replace('room_number', room_number)
                        if 'D1' in run.text:
                            run.text = run.text.replace('D1', difficult_airway_history)
                        if 'D2' in run.text:
                            run.text = run.text.replace('D2', physical_risk)
                        if 'R1' in run.text:
                            run.text = run.text.replace('R1', high_risk_desaturation)
                        if 'R2' in run.text:
                            run.text = run.text.replace('R2', high_risk_ICP)
                        if 'R3' in run.text:
                            run.text = run.text.replace('R3', unstable_hemodynamics)
                        if 'R4' in run.text:
                            run.text = run.text.replace('R4', other_risk_yes_no)
                        if 'risk_factors' in run.text:
                            run.text = run.text.replace('risk_factors', other_risk_text_input)
                        if 'who_will_intubate' in run.text:
                            if who_will_intubate:
                                run.text = run.text.replace('who_will_intubate', ', '.join(who_will_intubate).rstrip(', '))
                        if 'who_will_bvm' in run.text:
                            if who_will_bvm:
                                run.text = run.text.replace('who_will_bvm', ', '.join(who_will_bvm).rstrip(', '))
                        if 'intubation_method' in run.text:
                            run.text = run.text.replace('intubation_method', intubation_method)
                        if 'ett_type' in run.text:
                            run.text = run.text.replace('ett_type', ett_type)
                        if 'ett_size' in run.text:
                            run.text = run.text.replace('ett_size', ett_size)
                        if 'lma_details' in run.text:
                            run.text = run.text.replace('lma_details', lma_details)
                        if 'glide_details' in run.text:
                            run.text = run.text.replace('glide_details', glide_details)
                        if 'other_device_details' in run.text:
                            run.text = run.text.replace('other_device_details', other_device_details)
                        if 'mac_details' in run.text:
                            run.text = run.text.replace('mac_details', mac_details)
                        if 'miller_details' in run.text:
                            run.text = run.text.replace('miller_details', miller_details)
                        if 'wis_hipple_details' in run.text:
                            run.text = run.text.replace('wis_hipple_details', wis_hipple_details)
                        if 'atropine_dose' in run.text:
                            run.text = run.text.replace('atropine_dose', atropine_dose)
                        if 'glycopyrrolate_dose' in run.text:
                            run.text = run.text.replace('glycopyrrolate_dose', glycopyrrolate_dose)
                        if 'fentanyl_dose' in run.text:
                            run.text = run.text.replace('fentanyl_dose', fentanyl_dose)
                        if 'midazolam_dose' in run.text:
                            run.text = run.text.replace('midazolam_dose', midazolam_dose)
                        if 'ketamine_dose' in run.text:
                            run.text = run.text.replace('ketamine_dose', ketamine_dose)
                        if 'propofol_dose' in run.text:
                            run.text = run.text.replace('propofol_dose', propofol_dose)
                        if 'roc_dose' in run.text:
                            run.text = run.text.replace('roc_dose', roc_dose)
                        if 'vec_dose' in run.text:
                            run.text = run.text.replace('vec_dose', vec_dose)
                        if 'ao_details' in run.text:
                            run.text = run.text.replace('ao_details', ao_details)
                        if 'additional_notes' in run.text:
                            run.text = run.text.replace('additional_notes', other_planning)
                        if 'when_intubate' in run.text:
                            if when_intubate:
                                run.text = run.text.replace('when_intubate', ', '.join(when_intubate).rstrip(', '))  # Join with comma and space, then strip
                        if 'advance_airway_provider' in run.text:
                            if advance_airway_provider:
                                run.text = run.text.replace('advance_airway_provider', ', '.join(advance_airway_provider).rstrip(', '))  # Join with comma and space, then strip
                        if 'advance_airway_procedure' in run.text:
                            if advance_airway_procedure:
                                run.text = run.text.replace('advance_airway_procedure', ', '.join(advance_airway_procedure).rstrip(', '))  # Join with comma and space, then strip

                                        
    # Save the modified document
    doc_file = 'airway_bundle_form.docx'
    doc.save(doc_file)
    return doc_file

def reset_input(default_value, key):
    if key not in st.session_state:
        st.session_state[key] = default_value
    current_value = st.text_input("", key=key)
    if current_value != st.session_state[key]:
        st.session_state[key] = current_value
    return current_value

def initialize_firebase():
    global FIREBASE_COLLECTION_NAME
    FIREBASE_KEY_JSON = os.getenv('FIREBASE_KEY')
    FIREBASE_COLLECTION_NAME = os.getenv('FIREBASE_COLLECTION_NAME')
    
    if FIREBASE_KEY_JSON is None:
        raise ValueError("FIREBASE_KEY environment variable not set.")

    try:
        firebase_credentials = json.loads(FIREBASE_KEY_JSON)

        if not firebase_admin._apps:
            cred = credentials.Certificate(firebase_credentials)
            firebase_admin.initialize_app(cred)

        return firestore.client()
    except Exception as e:
        raise Exception(f"Error initializing Firebase: {e}")

db = initialize_firebase()

def update_ett_size():
    selected_age = st.session_state.age_select
    st.session_state.ett_size = age_to_ett_mapping.get(selected_age, '')

def fill_word_template(template_path, data):
    doc = Document(template_path)
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if f'{{{{{key}}}}}' in paragraph.text:
                paragraph.text = paragraph.text.replace(f'{{{{{key}}}}}', str(value))
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def next_section():
    if st.session_state.section < 8:
        st.session_state.section += 1
        save_data()

def prev_section():
    if st.session_state.section > 0:
        st.session_state.section -= 1
    
def save_data():
    data = {key: st.session_state.form_data.get(key, '') for key in st.session_state.form_data.keys()}
    db.collection('airway_checklists').add(data)
    
default_values = {
    'section': 0,
    'form_data': {},
    'selected_age': "",
    'selected_weight': "",
    'option': None,
    'completed_by': None,
    'room_number': None,
    'difficult_airway_history': 'Select Risk Factor 1',
    'physical_risk': 'Select Risk Factor 2',
    'high_risk_desaturation': 'Select Risk Factor 3',
    'high_risk_ICP': 'Select Risk Factor 4',
    'unstable_hemodynamics': 'Select Risk Factor 5',
    'other_risk_yes_no': 'Select Risk Factor 6',
    'other_risk_text_input': '',
    'who_will_intubate': [],  # Change to list if needed
    'who_will_bvm': [],       # Change to list if needed
    'intubation_method': "Intubation Method",
    'ett_size': None,
    'ett_type': None,
    'lma_details': None,
    'glide_details': None,
    'other_device_details': None,
    'mac_details': None,
    'miller_details': None,
    'wis_hipple_details': "",
    'atropine_dose': None,
    'glycopyrrolate_dose': None,
'fentanyl_dose': None,
'midazolam_dose': None,
'ketamine_dose': "",
'propofol_dose': "",
'roc_dose': None,
'vec_dose': None,
'ao_details': None, 
'other_planning': "",
'when_intubate': [],
'advance_airway_provider': [],
'advance_airway_procedure': [],
}

# Initialize session state variables if not already set
for key, value in default_values.items():
    if key not in st.session_state:
        st.session_state[key] = value

#Front Page Completed Section
if st.session_state['section'] == 0:
    st.title("NEARS4KIDS BUNDLE FORM")
    st.markdown("")
    st.markdown("Front Page")
    
    # Selectbox for front page completion
    option = st.selectbox(
        "Select an option", 
        [
            "Select an option", 
            "On admission", 
            "During rounds", 
            "After Rounds", 
            "Just prior to intubation", 
            "After intubation", 
            "Prior to Extubation"
        ],
        index=["Select an option", "On admission", "During rounds", "After Rounds", 
               "Just prior to intubation", "After intubation", "Prior to Extubation"].index(st.session_state['option']) if st.session_state['option'] else 0
    )
    
    completed_by = st.text_input("Who completed the form? (Name or Role)", value=st.session_state['completed_by'])
    
    room_number = st.selectbox(
        "Select Room Number", 
        [
            'Select Room Number', '4102', '4104', '4106', '4108', '4110', 
            '4112', '4114', '4116', '4201', '4203', 
            '4209', '4211', '4213', '4215', '4217', 
            '4219', '4221', '4223'
        ],
        index=['Select Room Number', '4102', '4104', '4106', '4108', '4110', 
               '4112', '4114', '4116', '4201', '4203', 
               '4209', '4211', '4213', '4215', '4217', 
               '4219', '4221', '4223'].index(st.session_state['room_number']) if st.session_state['room_number'] else 0
    )
        # Add the 'Previous' button to the first column

    cols = st.columns(2)
  
    with cols[0]:
        if st.button("Previous", on_click=prev_section):
            pass

    with cols[1]: 
      if st.button("Next"):
          if option != "Select an option" and room_number != "Select Room Number" and completed_by:
              st.session_state.option = option
              st.session_state.completed_by = completed_by
              st.session_state.room_number = room_number
              st.session_state.section += 1  # Increment the section
              st.rerun()  # Force a rerun to reflect changes immediately
          else:
              st.warning("Please select an option.")

# Patient Information Section
elif st.session_state.section == 1:
    st.title("Patient Information")

    cols = st.columns(2)
    
    # Set timezone to Eastern Time
    eastern = pytz.timezone('US/Eastern')
    
    with cols[0]:
        # Use a default value for date in EST
        now_est = datetime.now(eastern)
        date = st.date_input("Select Date (MM-DD-YYYY)", value=now_est.date(), key="date")

        if date:
            st.session_state['formatted_date'] = date.strftime("%m-%d-%Y")
            
        # Select Patient Age
        #age = st.selectbox("Select Patient Age", options=[""] + list(age_to_ett_mapping.keys()), key="age_select", on_change=update_automatic_selections)

        age_options = [""] + list(age_to_ett_mapping.keys())
        age_index = age_options.index(st.session_state.selected_age) if st.session_state.selected_age in age_options else 0

        age = st.selectbox(
            "Select Patient Age",
            options=age_options,
            index=age_index,
            key="age_select", on_change=update_automatic_selections
        )
      
        if age != st.session_state.selected_age:
            st.session_state.selected_age = age
        
    with cols[1]:
        # Use the current time in EST for time input
        current_time_est = datetime.now(eastern).time()
        time = st.time_input("Select Time", value=current_time_est, key="time")

        if time:
            st.session_state['formatted_time'] = time.strftime('%H:%M:%S')
            
        #weight = st.selectbox("Enter Patient Weight (Kilograms)", options=[""] + list(weight_to_atropine_mapping.keys()), key="weight_select", on_change=update_automatic_selections)

        weight_options = [""] + list(weight_to_atropine_mapping.keys())
        weight_index = weight_options.index(st.session_state.selected_weight) if st.session_state.selected_weight in weight_options else 0
        weight = st.selectbox(
            "Select Patient Weight (Kilograms)",
            options=weight_options,
            index=weight_index,
            key="weight_select", on_change=update_automatic_selections
        )

        if weight != st.session_state.selected_weight:
            st.session_state.selected_weight = weight
          
    # Initialize 'ett_size' in session state if it's not already set
    if 'ett_size' not in st.session_state:
        st.session_state['ett_size'] = ''  # Default value for ETT size
    
    selected_age = st.session_state.age_select
    
    st.session_state['ett_size'] = age_to_ett_mapping.get(selected_age, '')  # Update the session state with ETT size

    if 'glide_details' not in st.session_state:
        st.session_state['glide_details'] = ''  # Default value for ETT size

    st.session_state['glide_details'] = age_to_glide_mapping.get(selected_age, '')
    
    if 'mac_details' not in st.session_state:
        st.session_state['mac_details'] = ''  # Default value for ETT size
    
    st.session_state['mac_details'] = age_to_mac_mapping.get(selected_age, '')

    if 'miller_details' not in st.session_state:
        st.session_state['miller_details'] = ''  # Default value for ETT size
    
    st.session_state['miller_details'] = age_to_miller_mapping.get(selected_age, '')

    if 'ao_details' not in st.session_state:
        st.session_state['ao_details'] = ''  # Default value for ETT size
    
    st.session_state['ao_details'] = age_to_oxygenation_mapping.get(selected_age, '')

    if 'atropine_dose' not in st.session_state:
        st.session_state['atropine_dose'] = ''  # Default value for Atropine

    #selected_weight = st.session_state.weight_select
    
    if 'glycopyrrolate_dose' not in st.session_state:
        st.session_state['glycopyrrolate_dose'] = ''  # Default value for Glycopyrrolate
    
    if 'fentanyl_dose' not in st.session_state:
        st.session_state['fentanyl_dose'] = ''  # Default value for Fentanyl

    if 'lma_details' not in st.session_state:
        st.session_state['lma_details'] = ''  # Default value for ETT size

    
    # Retrieve the selected weight from session state
    selected_weight = st.session_state.get('weight_select', '')
    
    # If the weight is selected, update the drug doses accordingly (based on mappings)
    if selected_weight:
        st.session_state['lma_details'] = lma_size_mapping.get(selected_weight, '')
        st.session_state['atropine_dose'] = weight_to_atropine_mapping.get(selected_weight, '')
        st.session_state['glycopyrrolate_dose'] = weight_to_glycopyrrolate_mapping.get(selected_weight, '')
        st.session_state['fentanyl_dose'] = weight_to_fentanyl_mapping.get(selected_weight, '')

    # Default values for Midazolam, Ketamine, and Propofol if not set in session state
    if 'midazolam_dose' not in st.session_state:
        st.session_state['midazolam_dose'] = ''  # Default value for Midazolam
    
    if 'ketamine_dose' not in st.session_state:
        st.session_state['ketamine_dose'] = ''  # Default value for Ketamine
    
    if 'propofol_dose' not in st.session_state:
        st.session_state['propofol_dose'] = ''  # Default value for Propofol
    
    # Update doses based on the selected weight
    if selected_weight:
        st.session_state['midazolam_dose'] = weight_to_midaz_mapping.get(selected_weight, '')
        #st.session_state['ketamine_dose'] = weight_to_ketamine_mapping.get(selected_weight, '')
        #st.session_state['propofol_dose'] = weight_to_propo_mapping.get(selected_weight, '')

    # Default values for Rocuronium and Vecuronium if not set in session state
    if 'roc_dose' not in st.session_state:
        st.session_state['roc_dose'] = ''  # Default value for Rocuronium
    
    if 'vec_dose' not in st.session_state:
        st.session_state['vec_dose'] = ''  # Default value for Vecuronium
    
    # Update doses based on the selected weight
    if selected_weight:
        st.session_state['roc_dose'] = weight_to_roc_mapping.get(selected_weight, '')
        st.session_state['vec_dose'] = weight_to_vec_mapping.get(selected_weight, '')
    
    # Single Next and Previous Buttons
    col1, col2 = st.columns(2)

    # Add the 'Previous' button to the first column
    with col1:
        if st.button("Previous", on_click=prev_section):
            pass
          
    # Add the 'Next' button to the second column
    with col2:
        if st.button("Next", on_click=next_section):
            pass
            
# Intubation Risk Assessment Section
elif st.session_state.section == 2:
    st.title("Intubation Risk Assessment")
    st.write("#### Difficult Airway:")
    
    cols = st.columns([4, 1])
    with cols[0]:
        st.markdown("")
        st.markdown("")
        st.write("History of difficult airway?")
    with cols[1]:
        #difficult_airway_history = st.selectbox("", options=['Select Risk Factor 1', 'YES', 'NO'])
        difficult_airway_history = st.selectbox("", options=['Select Risk Factor 1', 'YES', 'NO'], index=['Select Risk Factor 1', 'YES', 'NO'].index(st.session_state.difficult_airway_history))

    cols = st.columns([4, 1])
    with cols[0]:
        st.markdown("")
        st.markdown("")
        st.write("Physical (e.g. small mouth, small jaw, large tongue, or short neck)?")
    
    with cols[1]:
        #physical_risk = st.selectbox("", options=['Select Risk Factor 2', 'YES', 'NO'])
        physical_risk = st.selectbox("", options=['Select Risk Factor 2', 'YES', 'NO'], index=['Select Risk Factor 2', 'YES', 'NO'].index(st.session_state.physical_risk))

    st.write("#### At Risk For:")
    
    cols = st.columns([4, 1])
    with cols[0]:
        st.markdown("")
        st.markdown("")
        st.write("High risk for rapid desaturation during intubation?")
    
    with cols[1]:
        #high_risk_desaturation = st.selectbox("", options=['Select Risk Factor 3', 'YES', 'NO'])
        high_risk_desaturation = st.selectbox("", options=['Select Risk Factor 3', 'YES', 'NO'], index=['Select Risk Factor 3', 'YES', 'NO'].index(st.session_state.high_risk_desaturation))

    cols = st.columns([4, 1])
    with cols[0]:
        st.markdown("")
        st.markdown("")
        st.write("Increased ICP, pulmonary hypertension, need to avoid hypercarbia?")
    
    with cols[1]:
        #high_risk_ICP = st.selectbox("", options=['Select Risk Factor 4', 'YES', 'NO'])
        high_risk_ICP = st.selectbox("", options=['Select Risk Factor 4', 'YES', 'NO'], index=['Select Risk Factor 4', 'YES', 'NO'].index(st.session_state.high_risk_ICP))

    cols = st.columns([4, 1])
    with cols[0]:
        st.markdown("")
        st.markdown("")
        st.write("Unstable hemodynamics (e.g., hypovolemia, potential need for fluid bolus, vasopressor, CPR)?")
    
    with cols[1]:
        #unstable_hemodynamics = st.selectbox("", options=['Select Risk Factor 5', 'YES', 'NO'])
        unstable_hemodynamics = st.selectbox("", options=['Select Risk Factor 5', 'YES', 'NO'], index=['Select Risk Factor 5', 'YES', 'NO'].index(st.session_state.unstable_hemodynamics))

    cols = st.columns([4, 1])
    
    # First column for the label
    with cols[0]:
        st.markdown("")
        st.markdown("")
        st.write("Other Risk Factors?")
    
    # Second column for the selectbox
    with cols[1]:
        #other_risk_yes_no = st.selectbox("", options=['Select Risk Factor 6', 'YES', 'NO'])
        other_risk_yes_no = st.selectbox("", options=['Select Risk Factor 6', 'YES', 'NO'], index=['Select Risk Factor 6', 'YES', 'NO'].index(st.session_state.other_risk_yes_no))

    with cols[0]:
        other_risk_text_input = ""
    
        if other_risk_yes_no == 'YES':
            #other_risk_text_input = st.text_input("Please specify the other risk:")
            other_risk_text_input = st.text_input("Please specify the other risk:", value=st.session_state.other_risk_text_input)

    # Single Next and Previous Buttons
    col1, col2, col3, col4, col5 = st.columns(5)

    # Add the 'Previous' button to the first column
    with col1:
        if st.button("Previous", on_click=prev_section):
            pass
    
    # Add the 'Next' button to the second column
    with col5:
        if st.button("Next"):
            if (difficult_airway_history != "Select Risk Factor 1" and 
                physical_risk != "Select Risk Factor 2" and 
                high_risk_desaturation != "Select Risk Factor 3" and 
                high_risk_ICP != "Select Risk Factor 4" and 
                unstable_hemodynamics != "Select Risk Factor 5" and 
                other_risk_yes_no != "Select Risk Factor 6"):
                
                # Include the other_risk_text_input in your logic
                st.session_state.difficult_airway_history = difficult_airway_history
                st.session_state.physical_risk = physical_risk
                st.session_state.high_risk_desaturation = high_risk_desaturation
                st.session_state.high_risk_ICP = high_risk_ICP
                st.session_state.unstable_hemodynamics = unstable_hemodynamics
                st.session_state.other_risk_yes_no = other_risk_yes_no
                
                if other_risk_yes_no == 'YES':
                    st.session_state.other_risk_text_input = other_risk_text_input
                else:
                    st.session_state.other_risk_text_input = ""  # or handle accordingly
                
                # Increment section and rerun
                st.session_state.section += 1
                st.rerun()
            else:
                st.warning("Please select all options.")

elif st.session_state.section == 3:
    st.title("Intubation Plan")

    who_will_intubate = st.multiselect(
    "Who will intubate?", 
    options=['Resident', 'Fellow', 'NP', 'Attending', 'Anesthesiologist', 'ENT physician', 'RT'],
    default=st.session_state.get('who_will_intubate', [])
    )
    
    who_will_bvm = st.multiselect(
        "Who will bag-mask?", 
        options=['Resident', 'Fellow', 'NP', 'Attending', 'RT'],
        default=st.session_state.get('who_will_bvm', [])
    )


    # Create a layout for intubation method
    #intubation_method = st.selectbox("How will we intubate? (Method)", ["Intubation Method", "Oral", "Nasal"].index(st.session_state.intubation_method))

    intubation_options = ["Oral", "Nasal"]

# Check if 'intubation_method' exists in session state and is a valid value
    if 'intubation_method' in st.session_state and st.session_state.intubation_method in intubation_options:
        # Use the index of the saved value
        intubation_method = st.selectbox(
            "How will we intubate? (Method)", 
            options=intubation_options,
            index=intubation_options.index(st.session_state.intubation_method)  # Set the index of the saved value
        )
    else:
        # Default to the first option ("Intubation Method")
        intubation_method = st.selectbox(
            "How will we intubate? (Method)", 
            options=intubation_options,
            index=0  # Default to the first option if no valid session state value exists
        )
    
    # Save the selected intubation method in the session state
    st.session_state.intubation_method = intubation_method

    # Create a layout for ETT Type and ETT Size
    cols = st.columns(2)

    with cols[0]:
        ett_type_options = ["Cuffed", "Uncuffed"]

        # Check if 'ett_type' exists in session state and is a valid value
        if 'ett_type' in st.session_state and st.session_state.ett_type in ett_type_options:
            # Use the index of the saved value
            ett_type = st.selectbox(
                "ETT Type", 
                options=ett_type_options,
                index=ett_type_options.index(st.session_state.ett_type)  # Set the index of the saved value
            )
        else:
            # Default to the first option ("Cuffed") if no valid session state value exists
            ett_type = st.selectbox(
                "ETT Type", 
                options=ett_type_options,
                index=0  # Default to the first option if no valid session state value exists
            )

        # Save the selected ETT type in the session state
        st.session_state.ett_type = ett_type

      
    with cols[1]:
        # Get the ETT sizes from a mapping or some source
        ett_sizes = list(set(age_to_ett_mapping.values()))  # Assuming age_to_ett_mapping is defined
        
        # Add a blank option (e.g., "Select an ETT Size" or "")
        ett_sizes = ["Select an ETT Size"] + ett_sizes  # Prepend a "Select" option

        custom_order = ["2.5 mm", "3.0 mm", "3.5 mm", "4.0 mm", "4.5 mm", "5.0 mm", "5.5 mm", "6.0 mm", "6.5 mm", "7.0 mm", "7.5 mm", "8.0 mm"] #new line

        ett_sizes = ["Select an ETT Size"] + [size for size in custom_order if size in ett_sizes] #new line

        # Check if 'ett_size' exists and is valid in session state
        if 'ett_size' in st.session_state and st.session_state['ett_size'] in ett_sizes:
            ett_size = st.selectbox(
                "ETT Size", 
                options=ett_sizes, 
                key="ett_size_display", 
                index=ett_sizes.index(st.session_state['ett_size'])
            )
        else:
            # Set a default value (e.g., first available size or blank if not found)
            ett_size = st.selectbox(
                "ETT Size", 
                options=ett_sizes, 
                key="ett_size_display",
                index=0  # Default to "Select an ETT Size"
            )
        
        # Save the selected ETT size in the session state
        st.session_state['ett_size'] = ett_size
    
    st.write("Device:")
    
    cols = st.columns(2)

    # Column 1: Dropdowns for "X" or empty
    #with cols[0]:
        # Dropdowns to choose if devices are selected or not (X = selected)
    #    device_1_selection = st.selectbox("Select Device", options=["", "X"], key="dropdown_1")
    #    device_2_selection = st.selectbox("Select Device", options=["", "X"], key="dropdown_2")
    #    device_3_selection = st.selectbox("Select Device", options=["", "X"], key="dropdown_3")
    #    device_4_selection = st.selectbox("Select Device", options=["", "X"], key="dropdown_4")
    
    # Column 2: Editable text inputs (reverts to the original value after the user moves away)
    with cols[0]:
        # These text inputs will reset to their default value if changed and the user moves away
        device_1_text = reset_input("Laryngoscope", key="laryngoscope_textx")
        device_2_text = reset_input("LMA", key="lma_textx")
        device_3_text = reset_input("Glidescope", key="glidescope_textx")
        device_4_text = reset_input("Other Device", key="other_device_textx")
    
    # Column 3: Additional details for each device (uneditable placeholders)
    with cols[1]:
        # Text Inputs with uneditable placeholders (details of each device)
        st.text_input("Laryngoscope details:", key="laryngoscope_details", disabled=False)
        #st.text_input("", key="laryngoscope_details", disabled=False)
      
        lma_details = list(set(lma_size_mapping.values()))  # Get unique ETT sizes
        lma_details = ["Select an LMA Size"] + lma_details
        custom_order = ['LMA Size 1', 'LMA Size 1.5', 'LMA Size 2', 'LMA Size 2.5', 'LMA Size 3', 'LMA Size 4', 'LMA Size 5']
        lma_details = ["Select an LMA Size"] + [size for size in custom_order if size in lma_details]
        if 'lma_details' in st.session_state and st.session_state['lma_details'] in lma_details:
          lma_details = st.selectbox(
            "LMA Size Details:",
            options = lma_details,
            key="lma_size_display",
            index=lma_details.index(st.session_state['lma_details'])
          )

        else: 
            lma_details = st.selectbox(
             "LMA Size Details:",
              options = lma_details,
              key = "lma_size_display",
             index = 0
            )
        
        #lma_details = st.selectbox("LMA Size Details:", options=lma_details, key="lma_size_display", index=lma_details.index(st.session_state['lma_details']) if st.session_state['lma_details'] in lma_details else 0)
        st.session_state['lma_details'] = lma_details
        
        glide_details = list(set(age_to_glide_mapping.values()))  # Get unique ETT sizes
        glide_details = ["Select a Glidescope Size"] + glide_details
        custom_order = ['Glidescope Size 1 or 2', 'Glidescope Size 2', 'Glidescope Size 3', 'Glidescope Size 3 or 4']
      
        glide_details = ["Select a Glidescope Size"] + [size for size in custom_order if size in glide_details]
        if 'glide_details' in st.session_state and st.session_state['glide_details'] in glide_details:
          glide_details = st.selectbox(
            "Glidescope Size Details:",
            options = glide_details,
            key="glide_size_display",
            index=glide_details.index(st.session_state['glide_details'])
          )

        else: 
            glide_details = st.selectbox(
              "Glidescope Size Details:",
              options = glide_details,
              key = "glide_size_display",
              index = 0
            )
        
        #glide_details = st.selectbox("Glidescope Details:", options=glide_details, key="glide_size_display", index=glide_details.index(st.session_state['glide_details']) if st.session_state['glide_details'] in glide_details else 0)
        st.session_state['glide_details'] = glide_details

        other_device_details = "" 
        other_device_details = st.text_input("Other Device details:", disabled=False)
    
    st.write("Blade:")
    
    cols = st.columns(2)

    # Column 1: Dropdowns for "X" or empty
    #with cols[0]:
        # Dropdowns to choose if devices are selected or not (X = selected)
        #blade_1_selection = st.selectbox("Select Device", options=["", "X"], key="dropdown_5")
        #blade_2_selection = st.selectbox("Select Device", options=["", "X"], key="dropdown_6")
        #blade_3_selection = st.selectbox("Select Device", options=["", "X"], key="dropdown_7")
        
    
    # Column 2: Editable text inputs (reverts to the original value after the user moves away)
    with cols[0]:
        # These text inputs will reset to their default value if changed and the user moves away
        blade_1_text = reset_input("Mac", key="macx")
        blade_2_text = reset_input("Miller", key="millerx")
        blade_3_text = reset_input("Wis-Hipple", key="wis_hipplex")
    
    # Column 3: Additional details for each device (uneditable placeholders)
    with cols[1]:
        mac_details = list(set(age_to_mac_mapping.values()))  # Get unique ETT sizes
        mac_details = ["Select a Mac Size"] + mac_details
        custom_order_mac = ['','Mac Blade Size 1', 'Mac Blade Size 2', 'Mac Blade Size 3', 'Mac Blade Size 4']
        mac_details = ["Select a Mac Size"] + [size for size in custom_order_mac if size in mac_details]
        
        if 'mac_details' in st.session_state and st.session_state['mac_details'] in mac_details:
          mac_details = st.selectbox(
              "Mac Details:",
              options=mac_details,
              key="mac_size_display",
              index=mac_details.index(st.session_state['mac_details'])
          )
        else: 
            mac_details = st.selectbox(
                "Mac Details:",
                options=mac_details,
                key="mac_size_display",
                index=0
            )

        #mac_details = st.selectbox("Mac Details:", options=mac_details, key="mac_size_display", index=mac_details.index(st.session_state['mac_details']) if st.session_state['mac_details'] in mac_details else 0)
        st.session_state['mac_details'] = mac_details

        miller_details = list(set(age_to_miller_mapping.values()))  # Get unique ETT sizes
        miller_details = ["Select a Miller Size"] + miller_details
        custom_order_miller = ['', 'Miller Blade Size 0', 'Miller Blade Size 1', 'Miller Blade Size 2', 'Miller Blade Size 3', 'Miller Blade Size 4']
        miller_details = ["Select a Miller Size"] + [size for size in custom_order_miller if size in miller_details]
        if 'miller_details' in st.session_state and st.session_state['miller_details'] in miller_details:
          miller_details = st.selectbox(
            "Miller Details:",
            options=miller_details,
            key="miller_size_display",
            index=miller_details.index(st.session_state['miller_details'])
            )
        else: 
            miller_details = st.selectbox(
                "Miller Details:",
                options=miller_details,
                key="miller_size_display",
                index=0
            )

        #miller_details = st.selectbox("Miller Details:", options=miller_details, key="miller_size_display", index=miller_details.index(st.session_state['miller_details']) if st.session_state['miller_details'] in miller_details else 0)
        st.session_state['miller_details'] = miller_details

        wis_hipple_details = st.text_input("Wis-Hipple Details:", disabled=False, value=st.session_state['wis_hipple_details'])
        

    st.write("Medications:")
    
    cols = st.columns(2)

    # Column 1: Dropdowns for "X" or empty
    #with cols[0]:
        # Dropdowns to choose if devices are selected or not (X = selected)
    #    med_1_selection = st.selectbox("Select Medication", options=["", "X"], key="dropdown_8")
    #    med_2_selection = st.selectbox("Select Medication", options=["", "X"], key="dropdown_9")
    #    med_3_selection = st.selectbox("Select Medication", options=["", "X"], key="dropdown_10")
    #    med_4_selection = st.selectbox("Select Medication", options=["", "X"], key="dropdown_11")
    #    med_5_selection = st.selectbox("Select Medication", options=["", "X"], key="dropdown_12")
    #    med_6_selection = st.selectbox("Select Medication", options=["", "X"], key="dropdown_13")
    #    med_7_selection = st.selectbox("Select Medication", options=["", "X"], key="dropdown_14")
    #    med_8_selection = st.selectbox("Select Medication", options=["", "X"], key="dropdown_15")
    
    # Column 2: Editable text inputs (reverts to the original value after the user moves away)
    with cols[0]:
        # These text inputs will reset to their default value if changed and the user moves away
        med_1_text = reset_input("Atropine", key="atropinex")
        med_2_text = reset_input("Glycopyrrolate", key="glycox")
        med_3_text = reset_input("Fentanyl", key="fentanylx")
        med_4_text = reset_input("Midazolam", key="midazolamx")
        med_5_text = reset_input("Ketamine", key="ketaminex")
        med_6_text = reset_input("Propofol", key="propofolx")
        med_7_text = reset_input("Rocuronium", key="rocx")
        med_8_text = reset_input("Vecuronium", key="vecx")

    # Column 3: Additional details for each device (uneditable placeholders)
    with cols[1]:
        # Text Inputs with uneditable placeholders (details of each device)
        #atropine_dose = list(set(weight_to_atropine_mapping.values()))  # Get unique Atropine doses
        #atropine_dose = st.selectbox("Atropine Dose:", options=atropine_dose, key="atropine_dose_display",index=atropine_dose.index(st.session_state['atropine_dose']) if st.session_state['atropine_dose'] in atropine_dose else 0)
        #st.session_state['atropine_dose'] = atropine_dose
      
        atropine_dose = list(set(weight_to_atropine_mapping.values()))  
        atropine_dose = ["Select an Atropine Dose"] + atropine_dose

        custom_order_atropine = [
            '0.1 mg', '0.11 mg', '0.12 mg', '0.13 mg', '0.14 mg', '0.15 mg',
            '0.16 mg', '0.17 mg', '0.18 mg', '0.19 mg', '0.2 mg', '0.22 mg',
            '0.24 mg', '0.26 mg', '0.28 mg', '0.3 mg', '0.32 mg', '0.34 mg',
            '0.36 mg', '0.38 mg', '0.4 mg', '0.42 mg', '0.44 mg', '0.46 mg',
            '0.48 mg', '0.5 mg', '0.52 mg', '0.54 mg', '0.56 mg', '0.58 mg',
            '0.6 mg', '0.62 mg', '0.64 mg', '0.66 mg', '0.68 mg', '0.7 mg',
            '0.72 mg', '0.74 mg', '0.76 mg', '0.78 mg', '0.8 mg', '0.82 mg',
            '0.84 mg', '0.86 mg', '0.88 mg', '0.9 mg', '0.92 mg', '0.94 mg',
            '0.96 mg', '0.98 mg', '1.0 mg'
        ]
        
        # Ensure custom order is respected
        atropine_dose = ["Select an Atropine Dose"] + [dose for dose in custom_order_atropine if dose in atropine_dose]
        
        # Check session state for existing atropine_dose
        if 'atropine_dose' in st.session_state and st.session_state['atropine_dose'] in atropine_dose:
            atropine_dose = st.selectbox(
                "Atropine Dose:",
                options=atropine_dose,
                key="atropine_dose_display",
                index=atropine_dose.index(st.session_state['atropine_dose'])
            )
        else: 
            atropine_dose = st.selectbox(
                "Atropine Dose:",
                options=atropine_dose,
                key="atropine_dose_display",
                index=0
            )
        
        # Update session state with selected atropine_dose
        st.session_state['atropine_dose'] = atropine_dose

        
        #glycopyrrolate_dose = list(set(weight_to_glycopyrrolate_mapping.values()))  # Get unique Glycopyrrolate doses
        #glycopyrrolate_dose = st.selectbox("Glycopyrrolate Dose:",options=glycopyrrolate_dose, key="glycopyrrolate_dose_display",index=glycopyrrolate_dose.index(st.session_state['glycopyrrolate_dose']) if st.session_state['glycopyrrolate_dose'] in glycopyrrolate_dose else 0)
        #st.session_state['glycopyrrolate_dose'] = glycopyrrolate_dose
        
# Get unique Glycopyrrolate doses
        glycopyrrolate_dose = list(set(weight_to_glycopyrrolate_mapping.values()))  
        glycopyrrolate_dose = ["Select a Glycopyrrolate Dose"] + glycopyrrolate_dose
        custom_order_glycopyrrolate = [
            '0.008 mg', '0.012 mg', '0.016 mg', '0.02 mg', '0.024 mg', '0.028 mg',
            '0.032 mg', '0.036 mg', '0.04 mg', '0.048 mg', '0.056 mg', '0.064 mg',
            '0.072 mg', '0.08 mg', '0.1 mg'
        ]

        # Ensure custom order is respected
        glycopyrrolate_dose = ["Select a Glycopyrrolate Dose"] + [dose for dose in custom_order_glycopyrrolate if dose in glycopyrrolate_dose]
        
        # Check session state for existing glycopyrrolate_dose
        if 'glycopyrrolate_dose' in st.session_state and st.session_state['glycopyrrolate_dose'] in glycopyrrolate_dose:
            glycopyrrolate_dose = st.selectbox(
                "Glycopyrrolate Dose:",
                options=glycopyrrolate_dose,
                key="glycopyrrolate_dose_display",
                index=glycopyrrolate_dose.index(st.session_state['glycopyrrolate_dose'])
            )
        else: 
            glycopyrrolate_dose = st.selectbox(
                "Glycopyrrolate Dose:",
                options=glycopyrrolate_dose,
                key="glycopyrrolate_dose_display",
                index=0
            )
        
        # Update session state with selected glycopyrrolate_dose
        st.session_state['glycopyrrolate_dose'] = glycopyrrolate_dose

        #fentanyl_dose = list(set(weight_to_fentanyl_mapping.values()))  # Get unique Fentanyl doses
        #fentanyl_dose = st.selectbox("Fentanyl Dose:", options=fentanyl_dose, key="fentanyl_dose_display",index=fentanyl_dose.index(st.session_state['fentanyl_dose']) if st.session_state['fentanyl_dose'] in fentanyl_dose else 0)
        #st.session_state['fentanyl_dose'] = fentanyl_dose

        # Get unique Fentanyl doses

        fentanyl_dose = list(set(weight_to_fentanyl_mapping.values()))  
        fentanyl_dose = ["Select a Fentanyl Dose"] + fentanyl_dose
        custom_order_fentanyl = [
            '2.0 mcg', '3.0 mcg', '4.0 mcg', '5.0 mcg', '6.0 mcg', '7.0 mcg',
            '8.0 mcg', '9.0 mcg', '10.0 mcg', '12.0 mcg', '14.0 mcg', '16.0 mcg',
            '18.0 mcg', '20.0 mcg', '25.0 mcg', '30.0 mcg', '35.0 mcg', '40.0 mcg',
            '45.0 mcg', '50.0 mcg'
        ]

        # Ensure custom order is respected
        fentanyl_dose = ["Select a Fentanyl Dose"] + [dose for dose in custom_order_fentanyl if dose in fentanyl_dose]
        
        # Check session state for existing fentanyl_dose
        if 'fentanyl_dose' in st.session_state and st.session_state['fentanyl_dose'] in fentanyl_dose:
            fentanyl_dose = st.selectbox(
                "Fentanyl Dose:",
                options=fentanyl_dose,
                key="fentanyl_dose_display",
                index=fentanyl_dose.index(st.session_state['fentanyl_dose'])
            )
        else: 
            fentanyl_dose = st.selectbox(
                "Fentanyl Dose:",
                options=fentanyl_dose,
                key="fentanyl_dose_display",
                index=0
            )
        
        # Update session state with selected fentanyl_dose
        st.session_state['fentanyl_dose'] = fentanyl_dose

        
        #midazolam_dose = list(set(weight_to_midaz_mapping.values()))  # Get unique Midazolam doses
        #midazolam_dose = st.selectbox("Midazolam Dose:", options=midazolam_dose, key="midazolam_dose_display",index=midazolam_dose.index(st.session_state['midazolam_dose']) if st.session_state['midazolam_dose'] in midazolam_dose else 0)
        #st.session_state['midazolam_dose'] = midazolam_dose

        midazolam_dose = list(set(weight_to_midaz_mapping.values()))  
        midazolam_dose = ["Select a Midazolam Dose"] + midazolam_dose
        custom_order_midazolam = [
            '0.1 mg', '0.15 mg', '0.2 mg', '0.25 mg', '0.3 mg', '0.35 mg',
            '0.4 mg', '0.45 mg', '0.5 mg', '0.6 mg', '0.7 mg', '0.8 mg',
            '0.9 mg', '1.0 mg', '1.25 mg', '1.5 mg', '1.75 mg', '2.0 mg'
        ]
        # Ensure custom order is respected
        midazolam_dose = ["Select a Midazolam Dose"] + [dose for dose in custom_order_midazolam if dose in midazolam_dose]
        
        # Check session state for existing midazolam_dose
        if 'midazolam_dose' in st.session_state and st.session_state['midazolam_dose'] in midazolam_dose:
            midazolam_dose = st.selectbox(
                "Midazolam Dose:",
                options=midazolam_dose,
                key="midazolam_dose_display",
                index=midazolam_dose.index(st.session_state['midazolam_dose'])
            )
        else: 
            midazolam_dose = st.selectbox(
                "Midazolam Dose:",
                options=midazolam_dose,
                key="midazolam_dose_display",
                index=0
            )
        
        # Update session state with selected midazolam_dose
        st.session_state['midazolam_dose'] = midazolam_dose

        
        #ketamine_dose = list(set(weight_to_ketamine_mapping.values()))  # Get unique Ketamine doses
        #ketamine_dose = st.selectbox("Ketamine Dose:", options=ketamine_dose, key="ketamine_dose_display",index=ketamine_dose.index(st.session_state['ketamine_dose']) if st.session_state['ketamine_dose'] in ketamine_dose else 0)
        
        ketamine_dose = st.text_input("Ketamine Dose:", disabled=False, value=st.session_state['ketamine_dose'])
        st.session_state['ketamine_dose'] = ketamine_dose
        
        #propofol_dose = list(set(weight_to_propo_mapping.values()))  # Get unique Propofol doses
        #propofol_dose = st.selectbox("Propofol Dose:", options=propofol_dose, key="propofol_dose_display",index=propofol_dose.index(st.session_state['propofol_dose']) if st.session_state['propofol_dose'] in propofol_dose else 0)
        propofol_dose = st.text_input("Propofol Dose:", disabled=False, value=st.session_state['propofol_dose'])
        st.session_state['propofol_dose'] = propofol_dose
        
        #roc_dose = list(set(weight_to_roc_mapping.values()))  # Get unique Rocuronium doses
        #roc_dose = st.selectbox("Rocuronium Dose:", options=roc_dose, key="roc_dose_display",index=roc_dose.index(st.session_state['roc_dose']) if st.session_state['roc_dose'] in roc_dose else 0)
        #st.session_state['roc_dose'] = roc_dose

        # Get unique Rocuronium doses
        roc_dose = list(set(weight_to_roc_mapping.values()))  
        roc_dose = ["Select a Rocuronium Dose"] + roc_dose
        custom_order_roc = [
            '2.0 mg', '2.5 mg', '3.0 mg', '3.5 mg', '4.0 mg', '4.5 mg',
            '5.0 mg', '5.5 mg', '6.0 mg', '6.5 mg', '7.0 mg', '7.5 mg',
            '8.0 mg', '8.5 mg', '9.0 mg', '9.5 mg', '10.0 mg', '11.0 mg',
            '12.0 mg', '13.0 mg', '14.0 mg', '15.0 mg', '16.0 mg', '17.0 mg',
            '18.0 mg', '19.0 mg', '20.0 mg', '21.0 mg', '22.0 mg', '23.0 mg',
            '24.0 mg', '25.0 mg', '26.0 mg', '27.0 mg', '28.0 mg', '29.0 mg',
            '30.0 mg', '31.0 mg', '32.0 mg', '33.0 mg', '34.0 mg', '35.0 mg',
            '36.0 mg', '37.0 mg', '38.0 mg', '39.0 mg', '40.0 mg', '41.0 mg',
            '42.0 mg', '43.0 mg', '44.0 mg', '45.0 mg', '46.0 mg', '47.0 mg',
            '48.0 mg', '49.0 mg', '50.0 mg'
        ]
        # Ensure custom order is respected
        roc_dose = ["Select a Rocuronium Dose"] + [dose for dose in custom_order_roc if dose in roc_dose]
        
        # Check session state for existing roc_dose
        if 'roc_dose' in st.session_state and st.session_state['roc_dose'] in roc_dose:
            roc_dose = st.selectbox(
                "Rocuronium Dose:",
                options=roc_dose,
                key="roc_dose_display",
                index=roc_dose.index(st.session_state['roc_dose'])
            )
        else: 
            roc_dose = st.selectbox(
                "Rocuronium Dose:",
                options=roc_dose,
                key="roc_dose_display",
                index=0
            )
        
        # Update session state with selected roc_dose
        st.session_state['roc_dose'] = roc_dose

        
        # Get unique Vecuronium doses
        vec_dose = list(set(weight_to_vec_mapping.values()))  
        vec_dose = ["Select a Vecuronium Dose"] + vec_dose
        custom_order_vec = [
            '0.2 mg', '0.3 mg', '0.4 mg', '0.5 mg', '0.6 mg', '0.7 mg',
            '0.8 mg', '0.9 mg', '1.0 mg', '1.2 mg', '1.4 mg', '1.6 mg',
            '1.8 mg', '2.0 mg', '2.5 mg', '3.0 mg', '3.5 mg', '4.0 mg',
            '4.5 mg', '10.0 mg'
        ]
        # Ensure custom order is respected
        vec_dose = ["Select a Vecuronium Dose"] + [dose for dose in custom_order_vec if dose in vec_dose]
        
        # Check session state for existing vec_dose
        if 'vec_dose' in st.session_state and st.session_state['vec_dose'] in vec_dose:
            vec_dose = st.selectbox(
                "Vecuronium Dose:",
                options=vec_dose,
                key="vec_dose_display",
                index=vec_dose.index(st.session_state['vec_dose'])
            )
        else: 
            vec_dose = st.selectbox(
                "Vecuronium Dose:",
                options=vec_dose,
                key="vec_dose_display",
                index=0
            )
        
        # Update session state with selected vec_dose
        st.session_state['vec_dose'] = vec_dose


    st.write("Apneic Oxygenation:")
    
    cols = st.columns(2)

    # Column 1: Dropdowns for "X" or empty
    #with cols[0]:
        # Dropdowns to choose if devices are selected or not (X = selected)
    #    ao_selection = st.selectbox("Select Use", options=["Select if AO to be utilized", "Yes", "No"])
    
    # Column 2: Editable text inputs (reverts to the original value after the user moves away)
    with cols[0]:
        # These text inputs will reset to their default value if changed and the user moves away
        ao_text = reset_input("Apneic Oxygenation", key="aox")
    
    # Column 3: Additional details for each device (uneditable placeholders)
    with cols[1]:
        # Text Inputs with uneditable placeholders (details of each device)
        #st.text_input("Apneic Oxygenation Details:", key="ao_details", disabled=False)

        ao_details = list(set(age_to_oxygenation_mapping.values()))  # Get unique ETT sizes
        ao_details = st.selectbox("Apneic Oxygenation:", options=ao_details, key="ao_details_display", index=ao_details.index(st.session_state['ao_details']) if st.session_state['ao_details'] in ao_details else 0)
        st.session_state['ao_details'] = ao_details

    other_planning = "" 
    other_planning = st.text_input("Other Intubation Planning Details:", disabled=False)
    
    # Single Next and Previous Buttons
    col1, col2, col3 = st.columns(3)

    # Add the 'Previous' button to the first column
    with col1:
        if st.button("Previous", on_click=prev_section):
            pass
    
    # Add the 'Next' button to the second column
    with col3:
        if st.button("Next"): #and lma_details and glide_details and mac_details and miller_details
            if (who_will_intubate and who_will_bvm and 
                ett_type and ett_size and atropine_dose and
                glycopyrrolate_dose and 
                fentanyl_dose and 
                midazolam_dose and 
                roc_dose and vec_dose and ao_details):
                st.session_state.who_will_intubate = who_will_intubate
                st.session_state.who_will_bvm = who_will_bvm
                st.session_state.intubation_method = intubation_method
                st.session_state.ett_type = ett_type  # Store ETT type
                st.session_state.ett_size = ett_size 
                st.session_state.lma_details = lma_details 
                st.session_state.glide_details = glide_details 
                st.session_state.other_device_details = other_device_details 
                st.session_state.mac_details = mac_details 
                st.session_state.miller_details = miller_details 
                st.session_state.wis_hipple_details = wis_hipple_details 
                st.session_state.atropine_dose = atropine_dose 
                st.session_state.glycopyrrolate_dose = glycopyrrolate_dose
                st.session_state.fentanyl_dose = fentanyl_dose
                st.session_state.midazolam_dose = midazolam_dose
                st.session_state.ketamine_dose = ketamine_dose
                st.session_state.propofol_dose = propofol_dose
                st.session_state.roc_dose = roc_dose
                st.session_state.vec_dose = vec_dose
                st.session_state.ao_details = ao_details
                st.session_state.other_planning = other_planning

                st.session_state.section += 1  # Increment the section
                st.rerun()  # Force a rerun to reflect changes immediately
            else:
                st.warning("Please select an option.")
                

if st.session_state.section == 4:
    st.title("Timing of Intubation")

    #when_intubate = st.multiselect(
    #"When will we intubate? (Describe timing of airway management", 
    #options = ['Prior to procedure', 'Mental Status Changes', 'Hypoxemia Refractory to CPAP: SPO2 < 92%','Hypoxemia Refractory to CPAP: SPO2 < 90%','Hypoxemia Refractory to CPAP: SPO2 < 88%','Hypoxemia Refractory to CPAP: SPO2 < 85%','Hypoxemia Refractory to CPAP: SPO2 < 80%','Hypoxemia Refractory to CPAP: SPO2 < 75%''Ventilation failure refractory to NIV', 'Loss of Airway Protection'],
    #default=st.session_state.get('when_intubate', [])
    #)

    hypoxemia_options = [
        'Hypoxemia Refractory to CPAP: SPO2 < 92%',
        'Hypoxemia Refractory to CPAP: SPO2 < 90%',
        'Hypoxemia Refractory to CPAP: SPO2 < 88%',
        'Hypoxemia Refractory to CPAP: SPO2 < 85%',
        'Hypoxemia Refractory to CPAP: SPO2 < 80%',
        'Hypoxemia Refractory to CPAP: SPO2 < 75%'
    ]
    
    # Full options including non-hypoxemia options
    options = ['Prior to procedure', 'Mental Status Changes'] + hypoxemia_options + ['Ventilation failure refractory to NIV', 'Loss of Airway Protection']
    
    # Get the current selection from session state
    selected_options = st.session_state.get('when_intubate', [])
    
    # Initialize an empty message
    error_message = ""
    
    # Check for hypoxemia selection logic
    hypoxemia_selected = [opt for opt in selected_options if opt in hypoxemia_options]
    
    if len(hypoxemia_selected) > 1:
        # If more than one hypoxemia option is selected, show an error message
        error_message = "Please select only one 'Hypoxemia Refractory to CPAP' option. Your selections will be cleared."
        # Clear hypoxemia selections
        selected_options = [opt for opt in selected_options if opt not in hypoxemia_options]
        
    # Update session state with current selections
    st.session_state.when_intubate = selected_options
    
    # Use multiselect with all options
    when_intubate = st.multiselect(
        "When will we intubate? (Describe timing of airway management)", 
        options=options,
        default=selected_options,
    )
    
    # Display error message if applicable
    if error_message:
        st.warning(error_message)
  

    # Single Next and Previous Buttons
    col1, col2, col3 = st.columns(3)

    # Add the 'Previous' button to the first column
    with col1:
        if st.button("Previous", on_click=prev_section):
            pass
    
    # Add the 'Next' button to the second column
    #with col3:
    #    if st.button("Next"):
    #        if when_intubate:
    #            st.session_state.when_intubate = when_intubate
    #            st.session_state.section += 1  # Increment the section
    #            st.rerun()  # Force a rerun to reflect changes immediately
    #        else:
    #            st.warning("Please select an option.")

    with col3:
        if st.button("Next"):
            # Check for valid selections
            hypoxemia_selected = [opt for opt in when_intubate if opt in hypoxemia_options]
    
            if len(hypoxemia_selected) > 1:
                st.warning("Please select only one 'Hypoxemia Refractory to CPAP' option.")
            elif not when_intubate:
                st.warning("Please select an option.")
            else:
                # Valid selection, proceed to the next section
                st.session_state.when_intubate = when_intubate
                st.session_state.section += 1  # Increment the section
                st.success("Proceeding to the next section...")
                st.rerun() 
              
if st.session_state.section == 5:
    st.title("Advanced Airway Provider and Procedure")

    advance_airway_provider = st.multiselect(
    "Backup Advanced Airway Provider:", 
    options = ['Attending', 'Anesthesia', 'ENT', 'Fellow'],
    default=st.session_state.get('advance_airway_provider', [])
    )

    advance_airway_procedure = st.multiselect(
    "Difficult Airway Alert Procedure:", 
    options = ['Difficult Airway Cart', 'Difficult Airway Emergency Page'],
    default=st.session_state.get('advance_airway_procedure', [])
    )

    # Single Next and Previous Buttons
    col1, col2, col3 = st.columns(3)

    # Add the 'Previous' button to the first column
    with col1:
        if st.button("Previous", on_click=prev_section):
            pass
    
    # Add the 'Next' button to the second column
    with col3:
        if st.button("Next"):
            if advance_airway_provider and advance_airway_procedure:
                st.session_state.advance_airway_provider = advance_airway_provider
                st.session_state.advance_airway_procedure = advance_airway_procedure
                
                st.session_state.section += 1  # Increment the section
                st.rerun()  # Force a rerun to reflect changes immediately
            else:
                st.warning("Please select an option.")


elif st.session_state.section == 6:
    st.title("Download ABC Form")
    
    col1, col2, col3 = st.columns(3)

    with col3: 
            if st.button("Submit"):
                template_path = 'airway_bundlez.docx'  # Ensure this is the correct path
                
                try:
                    data = {
                        'date': st.session_state.formatted_date,
                        'time': st.session_state.formatted_time,
                        'option': st.session_state.option,
                        'completed_by': st.session_state.completed_by,
                        'room_number': st.session_state.room_number,
                        'difficult_airway_history': st.session_state.difficult_airway_history,
                        'physical_risk': st.session_state.physical_risk,
                        'high_risk_desaturation': st.session_state.high_risk_desaturation,
                        'high_risk_ICP': st.session_state.high_risk_ICP,
                        'unstable_hemodynamics': st.session_state.unstable_hemodynamics,
                        'other_risk_yes_no': st.session_state.other_risk_yes_no,
                        'other_risk_text_input': st.session_state.other_risk_text_input,
                        'who_will_intubate': st.session_state.who_will_intubate,
                        'who_will_bvm': st.session_state.who_will_bvm,
                        'intubation_method': st.session_state.intubation_method,
                        'ett_size': st.session_state.ett_size,
                        'ett_type': st.session_state.ett_type,
                        'lma_details': st.session_state.lma_details,
                        'glide_details': st.session_state.glide_details,
                        'other_device_details': st.session_state.other_device_details,
                        'mac_details': st.session_state.mac_details,
                        'miller_details': st.session_state.miller_details,
                        'wis_hipple_details': st.session_state.wis_hipple_details,
                        'atropine_dose': st.session_state.atropine_dose,
                        'glycopyrrolate_dose': st.session_state.glycopyrrolate_dose,
                        'fentanyl_dose': st.session_state.fentanyl_dose,
                        'midazolam_dose': st.session_state.midazolam_dose,
                        'ketamine_dose': st.session_state.ketamine_dose,
                        'propofol_dose': st.session_state.propofol_dose,
                        'roc_dose': st.session_state.roc_dose,
                        'vec_dose': st.session_state.vec_dose,
                        'ao_details': st.session_state.ao_details,
                        'other_planning': st.session_state.other_planning,
                        'when_intubate': st.session_state.when_intubate,
                        'advance_airway_provider': st.session_state.advance_airway_provider,
                        'advance_airway_procedure': st.session_state.advance_airway_procedure
                    }
                    
                    doc_file = create_word_doc(template_path, data)
                    
                    st.success("Document created successfully!")
    
                    with open(doc_file, 'rb') as f:
                        st.download_button(
                            label="Download Word Document",
                            data=f,
                            file_name=doc_file.split("/")[-1],  # Use only the file name
                            mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                        )
                    os.remove(doc_file)  # Clean up the file after download
                except Exception as e:
                    st.error(f"An error occurred: {e}")
                    st.exception(e)  # This will print the stack trace for debugging
            
    with col1:
        if st.button("Previous", on_click=prev_section):
            pass

        
