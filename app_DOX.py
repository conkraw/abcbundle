import streamlit as st
from docx import Document
import os

def create_word_doc(template_path, date, time, option, intubation_method, who_will_intubate, other_planning, additional_notes):
    # Load the Word document template
    doc = Document(template_path)

    # Function to replace text in a run
    def replace_placeholder(run, placeholder, replacement):
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, replacement)

    # Check and replace text in paragraphs
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            replace_placeholder(run, 'DatePlaceholder', date)
            replace_placeholder(run, 'TimePlaceholder', time)
            replace_placeholder(run, 'FrontPagePlaceholder', option)
            replace_placeholder(run, 'intubation_method', intubation_method)
            replace_placeholder(run, 'who_will_intubate', ', '.join(who_will_intubate))
            replace_placeholder(run, 'other_planning', other_planning)
            replace_placeholder(run, 'additional_notes', additional_notes)

    # Check and replace text in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        replace_placeholder(run, 'DatePlaceholder', date)
                        replace_placeholder(run, 'TimePlaceholder', time)
                        replace_placeholder(run, 'FrontPagePlaceholder', option)
                        replace_placeholder(run, 'intubation_method', intubation_method)
                        replace_placeholder(run, 'who_will_intubate', ', '.join(who_will_intubate))
                        replace_placeholder(run, 'other_planning', other_planning)
                        replace_placeholder(run, 'additional_notes', additional_notes)

    # Save the modified document
    doc_file = 'airway_bundle_form.docx'
    doc.save(doc_file)
    return doc_file

# Streamlit app
st.title("Fill in Template Document")

# Initialize session state
if 'page' not in st.session_state:
    st.session_state.page = 'date'

# Date input page
if st.session_state.page == 'date':
    date = st.text_input("Enter your date")
    
    if st.button("Next"):
        if date:
            st.session_state.date = date
            st.session_state.page = 'time'  # Navigate to time input page
        else:
            st.warning("Please enter a date.")

# Time input page
elif st.session_state.page == 'time':
    time = st.text_input("Enter your time")
    
    if st.button("Next"):
        if time:
            st.session_state.time = time
            st.session_state.page = 'option'  # Navigate to option selection page
        else:
            st.warning("Please enter a time.")

# Option selection page
elif st.session_state.page == 'option':
    option = st.selectbox("Select an option", [
        "Select an option", 
        "On admission", 
        "During rounds", 
        "After Rounds", 
        "Just prior to intubation", 
        "After intubation", 
        "Prior to Extubation"
    ])

    if st.button("Next"):
        if option != "Select an option":
            st.session_state.option = option
            st.session_state.page = 'intubation_method'  # Navigate to intubation method selection page
        else:
            st.warning("Please select an option.")

# Intubation method selection page
elif st.session_state.page == 'intubation_method':
    intubation_method = st.selectbox("Select an intubation method", [
        "Select a method",
        "Endotracheal tube",
        "Laryngeal mask airway",
        "Bougie",
        "Other"
    ])

    if st.button("Next"):
        if intubation_method != "Select a method":
            st.session_state.intubation_method = intubation_method
            st.session_state.page = 'who_will_intubate'  # Navigate to who will intubate selection page
        else:
            st.warning("Please select an intubation method.")

# Who will intubate selection page
elif st.session_state.page == 'who_will_intubate':
    st.subheader("Who will intubate?")
    who_will_intubate = st.multiselect("Select the names", [
        "Doctor A",
        "Doctor B",
        "Doctor C",
        "Nurse A",
        "Nurse B"
    ])

    if st.button("Next"):
        if who_will_intubate:
            st.session_state.who_will_intubate = who_will_intubate
            st.session_state.page = 'other_planning'  # Navigate to other planning page
        else:
            st.warning("Please select at least one person.")

# Other planning page
elif st.session_state.page == 'other_planning':
    other_planning = st.text_input("Enter additional planning details")

    if st.button("Next"):
        if other_planning:
            st.session_state.other_planning = other_planning
            st.session_state.page = 'additional_notes'  # Navigate to additional notes page
        else:
            st.warning("Please enter additional planning details.")

# Additional notes page
elif st.session_state.page == 'additional_notes':
    additional_notes = st.text_area("Enter any additional notes")

    if st.button("Next"):
        if additional_notes:
            st.session_state.additional_notes = additional_notes
            st.session_state.page = 'download'  # Navigate to download page
        else:
            st.warning("Please enter additional notes.")

# Download page
elif st.session_state.page == 'download':
    template_path = 'airway_bundlez.docx'  # Ensure this is the correct path

    # Debugging output
    st.write(f"Using template: {template_path}")
    st.write(f"Date entered: {st.session_state.date}")
    st.write(f"Time entered: {st.session_state.time}")
    st.write(f"Option selected: {st.session_state.option}")
    st.write(f"Intubation method selected: {st.session_state.intubation_method}")
    st.write(f"Who will intubate: {', '.join(st.session_state.who_will_intubate)}")
    st.write(f"Additional planning details: {st.session_state.other_planning}")
    st.write(f"Additional notes: {st.session_state.additional_notes}")

    try:
        doc_file = create_word_doc(
            template_path, 
            st.session_state.date, 
            st.session_state.time, 
            st.session_state.option, 
            st.session_state.intubation_method, 
            st.session_state.who_will_intubate,
            st.session_state.other_planning,  # Pass other planning details
            st.session_state.additional_notes  # Pass additional notes
        )
        st.success("Document created successfully!")
        
        with open(doc_file, 'rb') as f:
            st.download_button(
                label="Download Word Document",
                data=f,
                file_name=doc_file,
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        os.remove(doc_file)  # Clean up the file after download
    except Exception as e:
        st.error(f"An error occurred: {e}")

    if st.button("Go Back"):
        st.session_state.page = 'additional_notes'  # Navigate back to additional notes page


