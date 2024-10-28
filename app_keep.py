import streamlit as st
from docx import Document
from io import BytesIO
from datetime import datetime
import firebase_admin
from firebase_admin import credentials, firestore
import os
import json
from PyPDF2 import PdfReader, PdfWriter
from datetime import datetime
import pdfrw
import io



# Define mappings for ETT size, Blade type, and Apneic Oxygenation based on patient age
age_to_ett_mapping = {
    "0-1": "3.0",
    "2-5": "3.5",
    "6-12": "4.0",
    "13-18": "4.5",
    "18+": "5.0"
}

age_to_lma_mapping = {
    "0-1": "LMA #0",
    "2-5": "LMA #1",
    "6-12": "LMA #2",
    "13-18": "LMA #3",
    "18+": "LMA #4"
}

age_to_glide_mapping = {
    "0-1": "Glidescope #0",
    "2-5": "Glidescope #1",
    "6-12": "Glidescope #2",
    "13-18": "Glidescope #3",
    "18+": "Glidescope #4"
}

# Define other mappings based on age
age_to_mac_mapping = {
    "0-1": "mac Blade #0",
    "2-5": "mac Blade #1",
    "6-12": "mac Blade #2",
    "13-18": "mac Blade #3",
    "18+": "mac Blade #4"
}

age_to_miller_mapping = {
    "0-1": "miller Blade #0",
    "2-5": "miller Blade #1",
    "6-12": "miller Blade #2",
    "13-18": "miller Blade #3",
    "18+": "miller Blade #4"
}

weight_to_atropine_mapping = {
    "1 kg": "a1 mg",
    "2 kg": "a2 mg",
    "3 kg": "a3 mg",
    "4 kg": "a4 mg",
    "5 kg": "a5 mg"
}

weight_to_glycopyrrolate_mapping = {
    "1 kg": "g1 mcg",
    "2 kg": "g2 mcg",
    "3 kg": "g3 mcg",
    "4 kg": "g4 mcg",
    "5 kg": "g5 mcg"
}

weight_to_fentanyl_mapping = {
    "1 kg": "f1 mcg",
    "2 kg": "f2 mcg",
    "3 kg": "f3 mcg",
    "4 kg": "f4 mcg",
    "5 kg": "f5 mcg"
}


weight_to_midaz_mapping = {
    "1 kg": "m1 mcg",
    "2 kg": "m2 mcg",
    "3 kg": "m3 mcg",
    "4 kg": "m4 mcg",
    "5 kg": "m5 mcg"
}

weight_to_ketamine_mapping = {
    "1 kg": "k1 mcg",
    "2 kg": "k2 mcg",
    "3 kg": "k3 mcg",
    "4 kg": "k4 mcg",
    "5 kg": "k5 mcg"
}

weight_to_propo_mapping = {
    "1 kg": "p1 mcg",
    "2 kg": "p2 mcg",
    "3 kg": "p3 mcg",
    "4 kg": "p4 mcg",
    "5 kg": "p5 mcg"
}

weight_to_roc_mapping = {
    "1 kg": "r1 mcg",
    "2 kg": "r2 mcg",
    "3 kg": "r3 mcg",
    "4 kg": "r4 mcg",
    "5 kg": "r5 mcg"
}

weight_to_vec_mapping = {
    "1 kg": "v1 mcg",
    "2 kg": "v2 mcg",
    "3 kg": "v3 mcg",
    "4 kg": "v4 mcg",
    "5 kg": "v5 mcg"
}


age_to_oxygenation_mapping = {
    "0-1": "5 L/min",
    "2-5": "5 L/min",
    "6-12": "6 L/min",
    "13-18": "6 L/min",
    "18+": "8 L/min"
}

# Define a function to automatically update the other settings when the age is selected
def update_automatic_selections():
    # Check if age is selected (you can keep this or modify it as needed)
    if "age_select" in st.session_state and st.session_state.age_select:
        selected_age = st.session_state.age_select
        st.session_state.ett_size = age_to_ett_mapping[selected_age]
        st.session_state.lma_details = age_to_lma_mapping[selected_age]
        st.session_state.glide_details = age_to_glide_mapping[selected_age]
        st.session_state.mac_details = age_to_mac_mapping[selected_age]
        st.session_state.miller_details = age_to_miller_mapping[selected_age]
        st.session_state.oxygenation = age_to_oxygenation_mapping[selected_age]

    # Check if weight is selected
    if "weight_select" in st.session_state and st.session_state.weight_select:
        selected_weight = st.session_state.weight_select
        # Update drug dosages based on the selected weight
        st.session_state.atropine_dose = weight_to_atropine_mapping[selected_weight]
        st.session_state.glycopyrrolate_dose = weight_to_glycopyrrolate_mapping[selected_weight]
        st.session_state.fentanyl_dose = weight_to_fentanyl_mapping[selected_weight]
        st.session_state.midazolam_dose = weight_to_midaz_mapping[selected_weight]
        st.session_state.ketamine_dose = weight_to_ketamine_mapping[selected_weight]
        st.session_state.propofol_dose = weight_to_propo_mapping[selected_weight]
        st.session_state.roc_dose = weight_to_roc_mapping[selected_weight]
        st.session_state.vec_dose = weight_to_vec_mapping[selected_weight]

def create_word_doc(template_path, data):
    # Load the Word document template
    doc = Document(template_path)
    
    # Access parameters
    date = data.get('date')
    time = data.get('time')
    option = data.get('option')
    completed_by = data.get('completed_by')
    room_number = data.get('room_number')
    difficult_airway_history = data.get('difficult_airway_history')
    physical_risk = data.get('physical_risk')
    high_risk_desaturation = data.get('high_risk_desaturation')
    high_risk_ICP = data.get('high_risk_ICP')
    unstable_hemodynamics = data.get('unstable_hemodynamics')
    other_risk_yes_no = data.get('other_risk_yes_no')
    other_risk_text_input = data.get('other_risk_text_input')
    who_will_intubate = data.get('who_will_intubate')
    who_will_bvm = data.get('who_will_bvm')
    other_intubate = data.get('other_intubate')
    other_bvm = data.get('other_bvm')
    intubation_method = data.get('intubation_method')

    # Check and replace text in paragraphs
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            # Replace Date and Time Placeholders
            if 'DatePlaceholder' in run.text:
                run.text = run.text.replace('DatePlaceholder', date)
            if 'TimePlaceholder' in run.text:
                run.text = run.text.replace('TimePlaceholder', time)
            if 'FrontPagePlaceholder' in run.text:
                run.text = run.text.replace('FrontPagePlaceholder', option)
            if 'DocumenterPlaceholder' in run.text:
                run.text = run.text.replace('DocumenterPlaceholder', completed_by)
            if 'room_number' in run.text:
                run.text = run.text.replace('room_number', room_number)
            if 'D1' in run.text:
                run.text = run.text.replace('D1', difficult_airway_history)
            if 'D2' in run.text:
                run.text = run.text.replace('D2', physical_risk)
            if 'R1' in run.text:
                run.text = run.text.replace('R1', high_risk_desaturation)
            if 'R2' in run.text:
                run.text = run.text.replace('R2', high_risk_ICP)
            if 'R3' in run.text:
                run.text = run.text.replace('R3', unstable_hemodynamics)
            if 'R4' in run.text:
                run.text = run.text.replace('R4', other_risk_yes_no)
            if 'risk_factors' in run.text:
                run.text = run.text.replace('risk_factors', other_risk_text_input)
            #if 'who_will_intubate' in run.text:
            #    run.text = run.text.replace('who_will_intubate', ' '.join(who_will_intubate))
            #if 'who_will_bvm' in run.text:
            #    run.text = run.text.replace('who_will_bvm', ' '.join(who_will_bvm))
            #if 'other_intubate' in run.text:
            #    run.text = run.text.replace('other_intubate', other_intubate)
            #if 'other_bvm' in run.text:
            #    run.text = run.text.replace('other_bvm', other_bvm)
            if 'intubation_method' in run.text:
                run.text = run.text.replace('intubation_method', intubation_method)

    # Save the modified document
    doc_file = 'airway_bundle_form.docx'
    doc.save(doc_file)
    return doc_file

def reset_input(default_value, key):
    if key not in st.session_state:
        st.session_state[key] = default_value
    current_value = st.text_input("", key=key)
    if current_value != st.session_state[key]:
        st.session_state[key] = current_value
    return current_value

def initialize_firebase():
    global FIREBASE_COLLECTION_NAME
    FIREBASE_KEY_JSON = os.getenv('FIREBASE_KEY')
    FIREBASE_COLLECTION_NAME = os.getenv('FIREBASE_COLLECTION_NAME')
    
    if FIREBASE_KEY_JSON is None:
        raise ValueError("FIREBASE_KEY environment variable not set.")

    try:
        firebase_credentials = json.loads(FIREBASE_KEY_JSON)

        if not firebase_admin._apps:
            cred = credentials.Certificate(firebase_credentials)
            firebase_admin.initialize_app(cred)

        return firestore.client()
    except Exception as e:
        raise Exception(f"Error initializing Firebase: {e}")

db = initialize_firebase()

def update_ett_size():
    selected_age = st.session_state.age_select
    st.session_state.ett_size = age_to_ett_mapping.get(selected_age, '')

def fill_word_template(template_path, data):
    doc = Document(template_path)
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if f'{{{{{key}}}}}' in paragraph.text:
                paragraph.text = paragraph.text.replace(f'{{{{{key}}}}}', str(value))
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def next_section():
    if st.session_state.section < 6:
        st.session_state.section += 1
        save_data()

def prev_section():
    if st.session_state.section > 0:
        st.session_state.section -= 1

def save_data():
    data = {key: st.session_state.form_data.get(key, '') for key in st.session_state.form_data.keys()}
    db.collection('airway_checklists').add(data)
    
default_values = {
    'section': 0,
    'form_data': {},
    'option': None,
    'completed_by': None,
    'room_number': None,
    'difficult_airway_history': 'Select Risk Factor 1',
    'physical_risk': 'Select Risk Factor 2',
    'high_risk_desaturation': 'Select Risk Factor 3',
    'high_risk_ICP': 'Select Risk Factor 4',
    'unstable_hemodynamics': 'Select Risk Factor 5',
    'other_risk_yes_no': 'Select Risk Factor 6',
    'other_risk_text_input': '',
    #'who_will_intubate': [],  # Change to list if needed
    #'who_will_bvm': [],       # Change to list if needed
    #'other_intubate': '',
    #'other_bvm': '',
    'intubation_method': None,
}

# Initialize session state variables if not already set
for key, value in default_values.items():
    if key not in st.session_state:
        st.session_state[key] = value
        
# Front Page Completed Section
if st.session_state.section == 0:
    st.title("Front Page Completed")
    
    # Selectbox for front page completion
    option = st.selectbox("Select an option", [
        "Select an option", 
        "On admission", 
        "During rounds", 
        "After Rounds", 
        "Just prior to intubation", 
        "After intubation", 
        "Prior to Extubation"
    ])
    
    completed_by = st.text_input("Who completed the form? (Name or Role)")
    
    room_number = st.selectbox("Select Room Number", 
                                ['Select Room Number', '4102', '4104', '4106', '4108', '4110', 
                                 '4112', '4114', '4116', '4201', '4203', 
                                 '4209', '4211', '4213', '4215', '4217', 
                                 '4219', '4221', '4223'])
    
    if st.button("Next"):
        if option != "Select an option" and room_number != "Select Room Number" and completed_by:
            st.session_state.option = option
            st.session_state.completed_by = completed_by
            st.session_state.room_number = room_number
            st.session_state.section += 1  # Increment the section
            st.rerun()  # Force a rerun to reflect changes immediately
        else:
            st.warning("Please select an option.")

# Patient Information Section
elif st.session_state.section == 1:
    st.title("Patient Information")

    cols = st.columns(2)
    
    with cols[0]:
    # Use a default value and store it in session state directly
        date = st.date_input("Select Date (MM-DD-YYYY)", value=datetime.today(), key="date")

        if date:
            st.session_state['formatted_date'] = date.strftime("%m-%d-%Y")
            
        # Select Patient Age
        age = st.selectbox("Select Patient Age",options=[""] + list(age_to_ett_mapping.keys()),key="age_select",on_change=update_automatic_selections)

    

    with cols[1]:
        time = st.time_input("Select Time", value=datetime.now().time(), key="time")

        if time:
            st.session_state['formatted_time'] = time.strftime('%H:%M:%S')
            
        weight = st.selectbox("Enter Patient Weight (Kilograms)", options=[""] + list(weight_to_atropine_mapping.keys()), key="weight_select",on_change=update_automatic_selections)

    # Initialize 'ett_size' in session state if it's not already set
    if 'ett_size' not in st.session_state:
        st.session_state['ett_size'] = ''  # Default value for ETT size
    
    selected_age = st.session_state.age_select
    
    st.session_state['ett_size'] = age_to_ett_mapping.get(selected_age, '')  # Update the session state with ETT size

    if 'lma_details' not in st.session_state:
        st.session_state['lma_details'] = ''  # Default value for ETT size
    
    st.session_state['lma_details'] = age_to_lma_mapping.get(selected_age, '')

    if 'glide_details' not in st.session_state:
        st.session_state['glide_details'] = ''  # Default value for ETT size

    st.session_state['glide_details'] = age_to_glide_mapping.get(selected_age, '')
    
    if 'mac_details' not in st.session_state:
        st.session_state['mac_details'] = ''  # Default value for ETT size
    
    st.session_state['mac_details'] = age_to_mac_mapping.get(selected_age, '')

    if 'miller_details' not in st.session_state:
        st.session_state['miller_details'] = ''  # Default value for ETT size
    
    st.session_state['miller_details'] = age_to_miller_mapping.get(selected_age, '')

    if 'ao_details' not in st.session_state:
        st.session_state['ao_details'] = ''  # Default value for ETT size
    
    st.session_state['ao_details'] = age_to_oxygenation_mapping.get(selected_age, '')

    if 'atropine_dose' not in st.session_state:
        st.session_state['atropine_dose'] = ''  # Default value for Atropine

    #selected_weight = st.session_state.weight_select
    
    if 'glycopyrrolate_dose' not in st.session_state:
        st.session_state['glycopyrrolate_dose'] = ''  # Default value for Glycopyrrolate
    
    if 'fentanyl_dose' not in st.session_state:
        st.session_state['fentanyl_dose'] = ''  # Default value for Fentanyl
    
    # Retrieve the selected weight from session state
    selected_weight = st.session_state.get('weight_select', '')
    
    # If the weight is selected, update the drug doses accordingly (based on mappings)
    if selected_weight:
        st.session_state['atropine_dose'] = weight_to_atropine_mapping.get(selected_weight, '')
        st.session_state['glycopyrrolate_dose'] = weight_to_glycopyrrolate_mapping.get(selected_weight, '')
        st.session_state['fentanyl_dose'] = weight_to_fentanyl_mapping.get(selected_weight, '')

    # Default values for Midazolam, Ketamine, and Propofol if not set in session state
    if 'midazolam_dose' not in st.session_state:
        st.session_state['midazolam_dose'] = ''  # Default value for Midazolam
    
    if 'ketamine_dose' not in st.session_state:
        st.session_state['ketamine_dose'] = ''  # Default value for Ketamine
    
    if 'propofol_dose' not in st.session_state:
        st.session_state['propofol_dose'] = ''  # Default value for Propofol
    
    # Update doses based on the selected weight
    if selected_weight:
        st.session_state['midazolam_dose'] = weight_to_midaz_mapping.get(selected_weight, '')
        st.session_state['ketamine_dose'] = weight_to_ketamine_mapping.get(selected_weight, '')
        st.session_state['propofol_dose'] = weight_to_propo_mapping.get(selected_weight, '')

    # Default values for Rocuronium and Vecuronium if not set in session state
    if 'roc_dose' not in st.session_state:
        st.session_state['roc_dose'] = ''  # Default value for Rocuronium
    
    if 'vec_dose' not in st.session_state:
        st.session_state['vec_dose'] = ''  # Default value for Vecuronium
    
    # Update doses based on the selected weight
    if selected_weight:
        st.session_state['roc_dose'] = weight_to_roc_mapping.get(selected_weight, '')
        st.session_state['vec_dose'] = weight_to_vec_mapping.get(selected_weight, '')
    
    # Single Next and Previous Buttons
    col1, col2 = st.columns(2)

    # Add the 'Previous' button to the first column
    with col1:
        if st.button("Previous", on_click=prev_section):
            pass
    
    # Add the 'Next' button to the second column
    with col2:
        if st.button("Next", on_click=next_section):
            pass
            
# Intubation Risk Assessment Section
elif st.session_state.section == 2:
    st.title("Intubation Risk Assessment")
    st.write("#### Difficult Airway:")
    
    cols = st.columns([4, 1])
    with cols[0]:
        st.markdown("")
        st.markdown("")
        st.write("History of difficult airway?")
    with cols[1]:
        difficult_airway_history = st.selectbox("", options=['Select Risk Factor 1', 'YES', 'NO'])

    cols = st.columns([4, 1])
    with cols[0]:
        st.markdown("")
        st.markdown("")
        st.write("Physical (e.g. small mouth, small jaw, large tongue, or short neck)?")
    
    with cols[1]:
        physical_risk = st.selectbox("", options=['Select Risk Factor 2', 'YES', 'NO'])

    st.write("#### At Risk For:")
    
    cols = st.columns([4, 1])
    with cols[0]:
        st.markdown("")
        st.markdown("")
        st.write("High risk for rapid desaturation during intubation?")
    
    with cols[1]:
        high_risk_desaturation = st.selectbox("", options=['Select Risk Factor 3', 'YES', 'NO'])

    cols = st.columns([4, 1])
    with cols[0]:
        st.markdown("")
        st.markdown("")
        st.write("Increased ICP, pulmonary hypertension, need to avoid hypercarbia?")
    
    with cols[1]:
        high_risk_ICP = st.selectbox("", options=['Select Risk Factor 4', 'YES', 'NO'])

    cols = st.columns([4, 1])
    with cols[0]:
        st.markdown("")
        st.markdown("")
        st.write("Unstable hemodynamics (e.g., hypovolemia, potential need for fluid bolus, vasopressor, CPR)?")
    
    with cols[1]:
        unstable_hemodynamics = st.selectbox("", options=['Select Risk Factor 5', 'YES', 'NO'])

    cols = st.columns([4, 1])
    
    # First column for the label
    with cols[0]:
        st.markdown("")
        st.markdown("")
        st.write("Other Risk Factors?")
    
    # Second column for the selectbox
    with cols[1]:
        other_risk_yes_no = st.selectbox("", options=['Select Risk Factor 6', 'YES', 'NO'])

    with cols[0]:
        other_risk_text_input = ""
    
        if other_risk_yes_no == 'YES':
            other_risk_text_input = st.text_input("Please specify the other risk:")

    # Single Next and Previous Buttons
    col1, col2, col3 = st.columns(3)

    # Add the 'Previous' button to the first column
    with col1:
        if st.button("Previous", on_click=prev_section):
            pass
    
    # Add the 'Next' button to the second column
    with col3:
        if st.button("Next"):
            if (difficult_airway_history != "Select Risk Factor 1" and 
                physical_risk != "Select Risk Factor 2" and 
                high_risk_desaturation != "Select Risk Factor 3" and 
                high_risk_ICP != "Select Risk Factor 4" and 
                unstable_hemodynamics != "Select Risk Factor 5" and 
                other_risk_yes_no != "Select Risk Factor 6"):
                
                # Include the other_risk_text_input in your logic
                st.session_state.difficult_airway_history = difficult_airway_history
                st.session_state.physical_risk = physical_risk
                st.session_state.high_risk_desaturation = high_risk_desaturation
                st.session_state.high_risk_ICP = high_risk_ICP
                st.session_state.unstable_hemodynamics = unstable_hemodynamics
                st.session_state.other_risk_yes_no = other_risk_yes_no
                
                if other_risk_yes_no == 'YES':
                    st.session_state.other_risk_text_input = other_risk_text_input
                else:
                    st.session_state.other_risk_text_input = ""  # or handle accordingly
                
                # Increment section and rerun
                st.session_state.section += 1
                st.rerun()
            else:
                st.warning("Please select all options.")
    
# Intubation Plan Section
elif st.session_state.section == 3:
    st.title("Intubation Plan")
    who_will_intubate = st.multiselect("Who will intubate?", 
                                   ['Select_Intubator','Resident', 'Fellow', 'NP', 'Attending','Anesthesiologist','ENT physician','RT','Other Intubator:'])

    other_intubate = ""
    
    if 'Other Intubator:' in who_will_intubate:
        other_intubate = st.text_input("Please specify the 'other' clinician who will intubate:")
    
    who_will_bvm = st.multiselect("Who will bag-mask?", 
                                   ['Select_BVMer','Resident', 'Fellow', 'NP', 'Attending', 'RT', 'Other BVMer:'])
    other_bvm = ""
    
    if 'Other BVMer:' in who_will_bvm:
        other_bvm = st.text_input("Please specify the 'other' clinician who will perform bag mask valve ventilation:")
        
    # Create a layout for intubation method
    intubation_method = st.selectbox("How will we intubate? (Method)", ["Intubation Method","Oral", "Nasal"])

    # Create a layout for ETT Type and ETT Size
    cols = st.columns(2)

    with cols[0]:
        ett_type = st.selectbox("ETT Type", ["Cuffed", "Uncuffed"], key="ett_type")

    with cols[1]:

        ett_sizes = list(set(age_to_ett_mapping.values()))  # Get unique ETT sizes
        selected_ett_size = st.selectbox("ETT Size", options=ett_sizes, key="ett_size_display", index=ett_sizes.index(st.session_state['ett_size']) if st.session_state['ett_size'] in ett_sizes else 0)
        st.session_state['ett_size'] = selected_ett_size
        
    st.write("Device:")
    
    cols = st.columns(3)

    # Column 1: Dropdowns for "X" or empty
    with cols[0]:
        # Dropdowns to choose if devices are selected or not (X = selected)
        device_1_selection = st.selectbox("Select Device", options=["", "X"], key="dropdown_1")
        device_2_selection = st.selectbox("Select Device", options=["", "X"], key="dropdown_2")
        device_3_selection = st.selectbox("Select Device", options=["", "X"], key="dropdown_3")
        device_4_selection = st.selectbox("Select Device", options=["", "X"], key="dropdown_4")
    
    # Column 2: Editable text inputs (reverts to the original value after the user moves away)
    with cols[1]:
        # These text inputs will reset to their default value if changed and the user moves away
        device_1_text = reset_input("Laryngoscope", key="laryngoscope_textx")
        device_2_text = reset_input("Glidescope", key="glidescope_textx")
        device_3_text = reset_input("LMA", key="lma_textx")
        device_4_text = reset_input("Other Device", key="other_device_textx")
    
    # Column 3: Additional details for each device (uneditable placeholders)
    with cols[2]:
        # Text Inputs with uneditable placeholders (details of each device)
        st.text_input("Laryngoscope details:", key="laryngoscope_details", disabled=False)
        
        glide_details = list(set(age_to_glide_mapping.values()))  # Get unique ETT sizes
        selected_glide_details = st.selectbox("Glidescope Details:", options=glide_details, key="glide_size_display", index=glide_details.index(st.session_state['glide_details']) if st.session_state['glide_details'] in glide_details else 0)
        st.session_state['glide_details'] = selected_glide_details
        
        lma_details = list(set(age_to_lma_mapping.values()))  # Get unique ETT sizes
        selected_lma_details = st.selectbox("LMA Details:", options=lma_details, key="lma_display", index=lma_details.index(st.session_state['lma_details']) if st.session_state['lma_details'] in lma_details else 0)
        st.session_state['lma_details'] = selected_lma_details
        
        st.text_input("Other Device details:", key="other_device_details", disabled=False)

    st.write("Blade:")
    
    cols = st.columns(3)

    # Column 1: Dropdowns for "X" or empty
    with cols[0]:
        # Dropdowns to choose if devices are selected or not (X = selected)
        blade_1_selection = st.selectbox("Select Device", options=["", "X"], key="dropdown_5")
        blade_2_selection = st.selectbox("Select Device", options=["", "X"], key="dropdown_6")
        blade_3_selection = st.selectbox("Select Device", options=["", "X"], key="dropdown_7")
        
    
    # Column 2: Editable text inputs (reverts to the original value after the user moves away)
    with cols[1]:
        # These text inputs will reset to their default value if changed and the user moves away
        blade_1_text = reset_input("Mac", key="macx")
        blade_2_text = reset_input("Miller", key="millerx")
        blade_3_text = reset_input("Wis-Hipple", key="wis_hipplex")
    
    # Column 3: Additional details for each device (uneditable placeholders)
    with cols[2]:
        mac_details = list(set(age_to_mac_mapping.values()))  # Get unique ETT sizes
        selected_mac_details = st.selectbox("Mac Details:", options=mac_details, key="mac_size_display", index=mac_details.index(st.session_state['mac_details']) if st.session_state['mac_details'] in mac_details else 0)
        st.session_state['mac_details'] = selected_mac_details

        miller_details = list(set(age_to_miller_mapping.values()))  # Get unique ETT sizes
        selected_miller_details = st.selectbox("Miller Details:", options=miller_details, key="miller_size_display", index=miller_details.index(st.session_state['miller_details']) if st.session_state['miller_details'] in miller_details else 0)
        st.session_state['miller_details'] = selected_miller_details

        st.text_input("Wis-Hipple Details:", key="wis_hipple_details", disabled=False)
        
    st.write("Medications:")
    
    cols = st.columns(3)

    # Column 1: Dropdowns for "X" or empty
    with cols[0]:
        # Dropdowns to choose if devices are selected or not (X = selected)
        med_1_selection = st.selectbox("Select Medication", options=["", "X"], key="dropdown_8")
        med_2_selection = st.selectbox("Select Medication", options=["", "X"], key="dropdown_9")
        med_3_selection = st.selectbox("Select Medication", options=["", "X"], key="dropdown_10")
        med_4_selection = st.selectbox("Select Medication", options=["", "X"], key="dropdown_11")
        med_5_selection = st.selectbox("Select Medication", options=["", "X"], key="dropdown_12")
        med_6_selection = st.selectbox("Select Medication", options=["", "X"], key="dropdown_13")
        med_7_selection = st.selectbox("Select Medication", options=["", "X"], key="dropdown_14")
        med_8_selection = st.selectbox("Select Medication", options=["", "X"], key="dropdown_15")
    
    # Column 2: Editable text inputs (reverts to the original value after the user moves away)
    with cols[1]:
        # These text inputs will reset to their default value if changed and the user moves away
        med_1_text = reset_input("Atropine", key="atropinex")
        med_2_text = reset_input("Glycopyrrolate", key="glycox")
        med_3_text = reset_input("Fentanyl", key="fentanylx")
        med_4_text = reset_input("Midazolam", key="midazolamx")
        med_5_text = reset_input("Ketamine", key="ketaminex")
        med_6_text = reset_input("Propofol", key="propofolx")
        med_7_text = reset_input("Rocuronium", key="rocx")
        med_8_text = reset_input("Vecuronium", key="vecx")

    # Column 3: Additional details for each device (uneditable placeholders)
    with cols[2]:
        # Text Inputs with uneditable placeholders (details of each device)
        atropine_doses = list(set(weight_to_atropine_mapping.values()))  # Get unique Atropine doses
        selected_atropine_dose = st.selectbox("Atropine Dose:", options=atropine_doses, key="atropine_dose_display",index=atropine_doses.index(st.session_state['atropine_dose']) if st.session_state['atropine_dose'] in atropine_doses else 0)
        st.session_state['atropine_dose'] = selected_atropine_dose
        
        glycopyrrolate_doses = list(set(weight_to_glycopyrrolate_mapping.values()))  # Get unique Glycopyrrolate doses
        selected_glycopyrrolate_dose = st.selectbox("Glycopyrrolate Dose:",options=glycopyrrolate_doses, key="glycopyrrolate_dose_display",index=glycopyrrolate_doses.index(st.session_state['glycopyrrolate_dose']) if st.session_state['glycopyrrolate_dose'] in glycopyrrolate_doses else 0)
        st.session_state['glycopyrrolate_dose'] = selected_glycopyrrolate_dose

        fentanyl_doses = list(set(weight_to_fentanyl_mapping.values()))  # Get unique Fentanyl doses
        selected_fentanyl_dose = st.selectbox("Fentanyl Dose:", options=fentanyl_doses, key="fentanyl_dose_display",index=fentanyl_doses.index(st.session_state['fentanyl_dose']) if st.session_state['fentanyl_dose'] in fentanyl_doses else 0)
        st.session_state['fentanyl_dose'] = selected_fentanyl_dose
        
        midazolam_doses = list(set(weight_to_midaz_mapping.values()))  # Get unique Midazolam doses
        selected_midazolam_dose = st.selectbox("Midazolam Dose:", options=midazolam_doses, key="midazolam_dose_display",index=midazolam_doses.index(st.session_state['midazolam_dose']) if st.session_state['midazolam_dose'] in midazolam_doses else 0)
        st.session_state['midazolam_dose'] = selected_midazolam_dose
        
        ketamine_doses = list(set(weight_to_ketamine_mapping.values()))  # Get unique Ketamine doses
        selected_ketamine_dose = st.selectbox("Ketamine Dose:", options=ketamine_doses, key="ketamine_dose_display",index=ketamine_doses.index(st.session_state['ketamine_dose']) if st.session_state['ketamine_dose'] in ketamine_doses else 0)
        st.session_state['ketamine_dose'] = selected_ketamine_dose
        
        propofol_doses = list(set(weight_to_propo_mapping.values()))  # Get unique Propofol doses
        selected_propofol_dose = st.selectbox("Propofol Dose:", options=propofol_doses, key="propofol_dose_display",index=propofol_doses.index(st.session_state['propofol_dose']) if st.session_state['propofol_dose'] in propofol_doses else 0)
        st.session_state['propofol_dose'] = selected_propofol_dose
        
        roc_doses = list(set(weight_to_roc_mapping.values()))  # Get unique Rocuronium doses
        selected_roc_dose = st.selectbox("Rocuronium Dose:", options=roc_doses, key="roc_dose_display",index=roc_doses.index(st.session_state['roc_dose']) if st.session_state['roc_dose'] in roc_doses else 0)
        st.session_state['roc_dose'] = selected_roc_dose
        
        vec_doses = list(set(weight_to_vec_mapping.values()))  # Get unique Vecuronium doses
        selected_vec_dose = st.selectbox("Vecuronium Dose:", options=vec_doses, key="vec_dose_display",index=vec_doses.index(st.session_state['vec_dose']) if st.session_state['vec_dose'] in vec_doses else 0)
        st.session_state['vec_dose'] = selected_vec_dose

    st.write("Apneic Oxygenation:")
    
    cols = st.columns(3)

    # Column 1: Dropdowns for "X" or empty
    with cols[0]:
        # Dropdowns to choose if devices are selected or not (X = selected)
        ao_selection = st.selectbox("Select Use", options=["", "Yes", "No"], key="dropdown_16")
    
    # Column 2: Editable text inputs (reverts to the original value after the user moves away)
    with cols[1]:
        # These text inputs will reset to their default value if changed and the user moves away
        ao_text = reset_input("Apneic Oxygenation", key="aox")
    
    # Column 3: Additional details for each device (uneditable placeholders)
    with cols[2]:
        # Text Inputs with uneditable placeholders (details of each device)
        #st.text_input("Apneic Oxygenation Details:", key="ao_details", disabled=False)

        ao_details = list(set(age_to_oxygenation_mapping.values()))  # Get unique ETT sizes
        selected_ao_details = st.selectbox("Apneic Oxygenation:", options=ao_details, key="ao_details_display", index=ao_details.index(st.session_state['ao_details']) if st.session_state['ao_details'] in ao_details else 0)
        st.session_state['ao_details'] = selected_ao_details
            
    other_planning = st.text_input("Other Intubation Planning Details:", key="other_planning")

    # Single Next and Previous Buttons
    col1, col2, col3 = st.columns(3)

    # Add the 'Previous' button to the first column
    with col1:
        if st.button("Previous", on_click=prev_section):
            pass
    
    # Add the 'Next' button to the second column
    with col3:
        if st.button("Next"):
            if intubation_method != "Intubation Method":
                #st.session_state.who_will_intubate = who_will_intubate
                #st.session_state.who_will_bvm = who_will_bvm
                st.session_state.intubation_method = intubation_method
                
                #if who_will_intubate == 'Other Intubator:':
                #    st.session_state.other_intubate = other_intubate
                #else:
                #    st.session_state.other_intubate = ""  # or handle accordingly
                
                #if who_will_bvm == 'Other BVMer:':
                #    st.session_state.other_bvm = other_bvm
                #else:
                #    st.session_state.other_bvm = ""  # or handle accordingly
                    
                st.session_state.section += 1  # Increment the section
                st.rerun()  # Force a rerun to reflect changes immediately
            else:
                st.warning("Please select an option.")
                
# Timing of Intubation Section
elif st.session_state.section == 4:
    st.title("Timing of Intubation")
    when_intubate = st.multiselect(
        "When will we intubate? (Describe timing of airway management):",
        ['Prior to procedure', 'Mental Status Changes', 
         'Hypoxemia Refractory to CPAP', 'Ventilation failure refractory to NIV', 
         'Loss of Airway Protection', 'Other'],
        key="when_intubate"
    )

    if 'Hypoxemia Refractory to CPAP' in when_intubate:
        st.text_input("If the patient has refractory hypoxemia refractory to CPAP, it will be defined as a SPO2 Level Less than:", key="hypoxemia")

    if 'Other' in when_intubate:
        st.text_input("Please state an 'other' reason for the timing of intubation:", key="other_when_intubate")
        
    # Single Next and Previous Buttons
    col1, col2, col3 = st.columns(3)

    # Add the 'Previous' button to the first column
    with col1:
        if st.button("Previous", on_click=prev_section):
            pass
    
    # Add the 'Next' button to the second column
    with col3:
        if st.button("Next", on_click=next_section):
            pass

elif st.session_state.section == 5:
    st.title("Backup")

    # Multi-select for Advance Airway Provider
    advance_airway_provider = st.multiselect("Advance Airway Provider:", 
                                   ['Attending', 'Anesthesia', 'ENT', 'Fellow', 'Other'],
                                   key="advance_airway_provider")
    
    advance_airway_procedure = st.multiselect("Difficult Airway Procedure:", 
                                   ['Difficult Airway Cart','Difficult Airway Emergency Page', 'Other'],
                                   key="difficult_airway")

    if 'Other' in advance_airway_provider:
        st.text_input("Please state an 'other' Advanced Airway Provider", key="other_advance_airway_provider")

    if 'Other' in advance_airway_procedure:
        st.text_input("Please state an 'other' protocol for Difficult Airway Protocol Initiation:", key="other_advance_airway_procedure")

        
    # Single Next and Previous Buttons
    col1, col2, col3 = st.columns(3)

    # Add the 'Previous' button to the first column
    with col1:
        if st.button("Previous", on_click=prev_section):
            pass
    
    # Add the 'Next' button to the second column
    with col3:
        if st.button("Next", on_click=next_section):
            pass

elif st.session_state.section == 6:
    st.title("Fill in Template Document")
    
    col1, col2, col3 = st.columns(3)

    with col3: 
            if st.button("Submit"):
                template_path = 'airway_bundlex.docx'  # Ensure this is the correct path
    
                # Debugging output
                st.write(f"Using template: {template_path}")
                st.write(f"Date entered: {st.session_state.formatted_date}")
                st.write(f"Time entered: {st.session_state.formatted_time}")
                st.write(f"Option selected: {st.session_state.option}")
                st.write(f"Completed by: {st.session_state.completed_by}")
                st.write(f"Room number: {st.session_state.room_number}")
                st.write(f"Difficult airway history: {st.session_state.difficult_airway_history}")
                st.write(f"Physical Risk: {st.session_state.physical_risk}")
                st.write(f"High Risk Desaturation: {st.session_state.high_risk_desaturation}")
                st.write(f"high_risk_ICP: {st.session_state.high_risk_ICP}")
                st.write(f"unstable_hemodynamics: {st.session_state.unstable_hemodynamics}")
                st.write(f"other_risk_yes_no: {st.session_state.other_risk_yes_no}")
                st.write(f"other_risk_text_input: {st.session_state.other_risk_text_input}")
                st.write(f"Who will Intubate: {st.session_state.who_will_intubate}")
                st.write(f"Who will BVM: {st.session_state.who_will_bvm}")
                st.write(f"Other Intubator: {st.session_state.other_intubate}")
                st.write(f"Other BVMer: {st.session_state.other_bvm}")
                st.write(f"Method: {st.session_state.intubation_method}")
                
                
                try:
                    data = {
                        'date': st.session_state.formatted_date,
                        'time': st.session_state.formatted_time,
                        'option': st.session_state.option,
                        'completed_by': st.session_state.completed_by,
                        'room_number': st.session_state.room_number,
                        'difficult_airway_history': st.session_state.difficult_airway_history,
                        'physical_risk': st.session_state.physical_risk,
                        'high_risk_desaturation': st.session_state.high_risk_desaturation,
                        'high_risk_ICP': st.session_state.high_risk_ICP,
                        'unstable_hemodynamics': st.session_state.unstable_hemodynamics,
                        'other_risk_yes_no': st.session_state.other_risk_yes_no,
                        'other_risk_text_input': st.session_state.other_risk_text_input,
                        #'who_will_intubate': st.session_state.who_will_intubate,
                        #'who_will_bvm': st.session_state.who_will_bvm,
                        #'other_intubate': st.session_state.other_intubate,
                        #'other_bvm': st.session_state.other_bvm,
                        'intubation_method': st.session_state.intubation_method
                    }
                    doc_file = create_word_doc(template_path, data)
                    
                    st.success("Document created successfully!")
    
                    with open(doc_file, 'rb') as f:
                        st.download_button(
                            label="Download Word Document",
                            data=f,
                            file_name=doc_file.split("/")[-1],  # Use only the file name
                            mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                        )
                    os.remove(doc_file)  # Clean up the file after download
                except Exception as e:
                    st.error(f"An error occurred: {e}")
                    st.exception(e)  # This will print the stack trace for debugging
            
    with col1:
        if st.button("Previous", on_click=prev_section):
            pass
        
